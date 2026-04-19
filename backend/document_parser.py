"""
document_parser.py

PRIMARY:  Groq LLM (llama-3.3-70b-versatile)
          Split into THREE small focused calls to stay within token budget:
            Call 1 – Metadata (cover page text ~2500 chars)
            Call 2 – Synopsis / Criteria (targeted section text ~4000 chars)
            Call 3 – SoA table (raw table pipe-text only)
          + Python structural extraction for ALL sections/subsections (no token cost)

FALLBACK: Regex heuristics if Groq is unavailable or fails
"""

import re
import io
import os
import json
import html as html_module
import uuid
import logging
from typing import Optional, List, Dict, Any, Tuple

from dotenv import load_dotenv
load_dotenv(os.path.join(os.path.dirname(__file__), '.env'))

logger = logging.getLogger(__name__)

GROQ_API_KEY = os.getenv("GROQ_API_KEY", "").strip()
GROQ_MODEL   = "llama-3.1-8b-instant"

# ────────────────────────────────────────────────
# Standard protocol heading → section number map
# Standard protocol heading → section number map
# ────────────────────────────────────────────────
SECTION_HEADING_MAP = {
    # Front-matter unnumbered items
    'APPROVAL': 'approval',
    'VERSION OF PROTOCOL': 'version',
    # Section 1
    'INTRODUCTION': '1', 'BACKGROUND': '1', 'RATIONALE': '1',
    # Section 2
    'TRIAL OBJECTIVES': '2', 'OBJECTIVES AND ENDPOINTS': '2', 'OBJECTIVES': '2', 'ENDPOINTS': '2',
    # Section 3
    'INVESTIGATIONAL PLAN': '3', 'STUDY DESIGN': '3', 'OVERALL DESIGN': '3',
    # Section 4
    'SELECTION OF TRIAL POPULATION': '4', 'STUDY POPULATION': '4', 'POPULATION': '4', 'ELIGIBILITY': '4',
    # Section 5
    'TREATMENTS': '5', 'STUDY INTERVENTION': '5', 'INTERVENTION': '5', 'TREATMENT': '5', 'INVESTIGATIONAL MEDICINAL': '5',
    # Section 6
    'TRIAL PROCEDURES': '6', 'STUDY PROCEDURES': '6',
    # Section 7
    'TRIAL ASSESSMENTS': '7', 'STUDY ASSESSMENTS': '7', 'ASSESSMENTS AND PROCEDURES': '7',
    # Section 8
    'ADVERSE EVENTS': '8', 'ADVERSE EVENT': '8',
    # Section 9
    'STATISTICAL METHODS': '9', 'STATISTICAL': '9', 'STATISTICS': '9',
    # Section 10
    'DATA HANDLING': '10', 'DATA MANAGEMENT': '10',
    # Section 11
    'MONITORING': '11', 'MONITORING PROCEDURES': '11',
    # Section 12
    'CHANGES IN THE CONDUCT': '12', 'PROTOCOL AMENDMENTS': '12',
    # Section 13
    'REPORTING AND PUBLICATION': '13', 'REPORTING': '13',
    # Section 14
    'ETHICAL AND REGULATORY': '14', 'ETHICAL': '14', 'REGULATORY': '14',
    # Section 15
    'LIABILITIES AND INSURANCE': '15', 'LIABILITIES': '15',
    # Section 16
    'ARCHIVING': '16',
    # Section 17
    'REFERENCES': '17', 'BIBLIOGRAPHY': '17',
    # Unnumbered
    'PROTOCOL SUMMARY': 'synopsis', 'SYNOPSIS': 'synopsis', 'SUMMARY': 'synopsis',
    'WITHDRAWAL': 'withdrawal', 'DISCONTINUATION': 'withdrawal',
    'SUPPORTING DOCUMENTATION': 'supporting',
    'APPENDICES': 'appendices',
}

# Titles may be long on one PDF line; cap generously so headings are not missed.
SECTION_NUMBER_RE = re.compile(r'^(\d{1,2})[\.\)\s]+(.{3,300})$')  # "1. Introduction"
SUBSECTION_RE     = re.compile(
    r'^(\d{1,2}\.\d{1,2}(?:\.\d{1,2})*)[\.\)\s]+(.{2,300})$'
)  # "1.1 Background" / "8.3.4 Long subsection title" (short lines)
# Same pattern without title length cap — PDFs often merge heading + paragraph on one line
SUBSECTION_HEAD_RE = re.compile(
    r'^(\d{1,2}\.\d{1,2}(?:\.\d{1,2})*)[\.\)\s]+(.+)$'
)

# NIH / FDA template noise that PyMuPDF often injects into body text
_TEMPLATE_BOILERPLATE_RE = re.compile(
    r'(?i)\bNIH[-\s]?FDA\s+Clinical\s+Trial\s+Protocol\s+Template\b[^\n]*|'
    r'\bVersion\s+Protocol\s*<#>\s*|\bClinical\s+Trial\s+Protocol\s+Template\b[^\n]*',
    re.MULTILINE,
)

# ────────────────────────────────────────────────
# GROQ PROMPT 1 — Metadata (scalar fields only)
# ────────────────────────────────────────────────
META_SYSTEM = (
    "You are a clinical trial protocol data extractor. "
    "Return ONLY a single valid JSON object. No markdown, no fences, no explanation. "
    "Only extract what is explicitly present in the text."
)

META_PROMPT = """\
Extract these fields from the clinical protocol text and return a JSON object:

{{
  "protocol_title": "",
  "protocol_number": "",
  "nct_number": "",
  "principal_investigator": "",
  "sponsor": "",
  "funded_by": "",
  "version_number": "",
  "protocol_date": "",
  "clinical_phase": "",
  "indication": "",
  "imp": "",
  "coordinating_investigator": "",
  "expert_committee": "",
  "sponsor_name_address": "",
  "trial_sites": "",
  "planned_period": "",
  "fpfv": "",
  "lplv": "",
  "num_patients": ""
}}

If a field is absent, use "". Do NOT guess.

TEXT:
---
{text}
---

Return ONLY the JSON."""

# ────────────────────────────────────────────────
# GROQ PROMPT 2 — Synopsis (lists: objectives, endpoints, criteria)
# ────────────────────────────────────────────────
SYN_SYSTEM = (
    "You are a clinical trial synopsis extractor. "
    "Return ONLY a single valid JSON object. No markdown, no fences, no explanation."
)

SYN_PROMPT = """\
Extract the following fields from the protocol text. Return them as a JSON object.
All list fields must be arrays of strings (one item per element, empty list [] if absent).

{{
  "primary_objectives": [],
  "secondary_objectives": [],
  "exploratory_objectives": [],
  "primary_endpoints": [],
  "secondary_endpoints": [],
  "exploratory_endpoints": [],
  "inclusion_criteria": [],
  "exclusion_criteria": [],
  "statistical_methods": ""
}}

Include endpoints described as co-primary, key secondary, main secondary, or under
Objectives and Endpoints / efficacy endpoints. Each distinct endpoint must be its own string.

TEXT:
---
{text}
---

Return ONLY the JSON."""

# ────────────────────────────────────────────────
# GROQ PROMPT 3 — Schedule of Activities table
# ────────────────────────────────────────────────
SOA_SYSTEM = (
    "You are a clinical trial Schedule of Activities extractor. "
    "Return ONLY a single valid JSON object. No markdown, no fences, no explanation."
)

SOA_PROMPT = """\
Extract the Schedule of Activities (SoA) table from the clinical trial tables below.

Return:
{{
  "headers": ["Procedure", "Visit1", "Visit2", ...],
  "rows": [
    ["Procedure name", "1", "0", "1", ...],
    ["Another procedure", "0", "1", "0", ...]
  ]
}}

RULES:
- headers[0] MUST be the literal string "Procedure"
- headers[1..n] = visit/timepoint names from the document (e.g. "Screening", "Day 1", "Week 4")
- rows[i][0] = procedure/assessment name
- rows[i][1..n] = "1" if performed at that visit (mark ✓ ✔ X x Y Yes • or any check symbol), else "0"
- Include ALL rows including sub-category rows
- CRITICAL: You must aggressively align the marks to the exact visit columns. Even if the raw text is jagged, trace which visit an "X" or mark corresponds to.
- CRITICAL: The length of EVERY row array `rows[i]` MUST EXACTLY EQUAL the length of the `headers` array! Fill trailing empty visits with "0".
- If no SoA table found, return {{"headers": [], "rows": []}}

TABLE DATA:
---
{tables}
---

Return ONLY the JSON."""

# ────────────────────────────────────────────────
# GROQ PROMPT 4 — Table of Contents (nested tree)
# ────────────────────────────────────────────────
TOC_SYSTEM = """You are a clinical protocol document parser.
Your job is to extract the Table of Contents from raw PDF text and return it as structured JSON.

RULES:
1. Return ONLY valid JSON — no explanation, no markdown, no code fences.
2. The JSON must be an array of section objects at the top level.
3. Each section object has exactly these fields:
   - "number": the section number as a string (e.g. "2.3.1") — empty string "" if unnumbered
   - "title": the section title WITHOUT the number prefix
   - "level": integer depth starting at 0 (0 = main section, 1 = subsection, 2 = sub-subsection, etc.)
   - "page": integer page number if found, else null
   - "children": array of child section objects (same structure, can be empty [])

4. Infer nesting from the numbering pattern: "1" is level 0, "1.1" is level 1, "1.1.1" is level 2, etc.
   For unnumbered entries (like ABBREVIATIONS, SYNOPSIS), treat as level 0.

5. FIX broken OCR text automatically — e.g. "B ackground" → "Background", "I ntroduction" → "Introduction",
   "Ex clusion" → "Exclusion", "P opulation" → "Population". Reconstruct the correct word.

6. SKIP entries that are clearly page headers, document metadata, version numbers,
   dot leaders (.........), standalone page numbers, or repeated "Table of contents" title lines.

7. DO NOT include "List of tables", "List of figures", "List of abbreviations", "Glossary of terms"
   as TOC sections — these are front-matter, skip them.

8. Sections like "ABBREVIATIONS", "SYNOPSIS", "PROTOCOL SUMMARY", "STATEMENT OF COMPLIANCE",
   "REFERENCES", "APPENDICES" are valid — include them as level 0 with number "".

9. Table entries like "Table 6-1 ...", "Figure 3-1 ..." that appear inside the TOC —
   include them as children of their parent section if present in the source text.

10. The output must capture ALL nesting levels exactly as they appear — do not flatten or skip any level.
"""

TOC_USER_TEMPLATE = """Here is the raw Table of Contents text extracted from a clinical protocol PDF.
Parse it and return the structured JSON as described.

RAW TOC TEXT:
---
{raw_toc}
---

Return only the JSON array. Start with [ and end with ]."""



# ════════════════════════════════════════════════
# DOCUMENT PARSER
# ════════════════════════════════════════════════
def _strip_template_noise(text: str) -> str:
    """Remove repeated protocol-template headers/footers from a text fragment."""
    if not text:
        return text
    t = _TEMPLATE_BOILERPLATE_RE.sub(' ', text)
    t = re.sub(r'\s{2,}', ' ', t).strip()
    return t


class DocumentParser:

    def parse(self, file_content: bytes, filename: str, upload_dir: str = "uploads") -> dict:
        ext = os.path.splitext(filename)[1].lower()
        if ext in ['.docx', '.doc']:
            return self._parse_docx(file_content, filename, upload_dir)
        elif ext == '.pdf':
            return self._parse_pdf(file_content, filename, upload_dir)
        else:
            raise ValueError(f"Unsupported file type '{ext}'. Upload a .docx or .pdf file.")

    # ──────────────────────────────────────────────
    # FILE READERS
    # ──────────────────────────────────────────────
    def _parse_docx(self, content: bytes, filename: str, upload_dir: str) -> dict:
        from docx import Document
        import zipfile

        doc = Document(io.BytesIO(content))
        paragraphs: List[Dict] = []
        tables_data: List[List[List[str]]] = []
        soa_image_url: Optional[str] = None

        for para in doc.paragraphs:
            text = _strip_template_noise(para.text.strip())
            if not text:
                continue
            is_bold = any(r.bold for r in para.runs if r.bold is not None)
            paragraphs.append({
                'text': text,
                'style': para.style.name if para.style else 'Normal',
                'bold': is_bold,
                'font_size': None
            })

        for table in doc.tables:
            tbl = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if any(any(c for c in row) for row in tbl):
                tables_data.append(tbl)

        try:
            with zipfile.ZipFile(io.BytesIO(content)) as z:
                for name in sorted(z.namelist()):
                    if name.startswith('word/media/') and any(
                        name.lower().endswith(e) for e in ['.png', '.jpg', '.jpeg', '.bmp', '.gif']
                    ):
                        img_bytes = z.read(name)
                        ext_t = name.split('.')[-1].lower()
                        soa_image_url = self._save_image(img_bytes, ext_t, upload_dir)
                        break
        except Exception as e:
            logger.warning(f"DOCX image extraction: {e}")

        full_text = '\n'.join(p['text'] for p in paragraphs)
        return self._run(paragraphs, full_text, tables_data, soa_image_url)

    def _parse_pdf(self, content: bytes, filename: str, upload_dir: str) -> dict:
        import fitz
        from pypdf import PdfReader

        doc = fitz.open(stream=content, filetype="pdf")
        paragraphs: List[Dict] = []
        tables_data: List[List[List[str]]] = []
        soa_image_url: Optional[str] = None

        for page in doc:
            prev_text_block = False
            for block in page.get_text("dict")["blocks"]:
                if block["type"] == 0:
                    # Paragraph break between distinct text blocks (preserves spacing vs one giant line)
                    if prev_text_block and paragraphs and (paragraphs[-1]['text'] or '').strip():
                        paragraphs.append({
                            'text': '',
                            'style': 'Normal',
                            'bold': False,
                            'font_size': 12
                        })
                    prev_text_block = True
                    for line in block["lines"]:
                        spans = line["spans"]
                        text = " ".join(s["text"] for s in spans).strip()
                        text = _strip_template_noise(text)
                        if not text:
                            continue
                        max_size = max((s["size"] for s in spans), default=12)
                        is_bold = any(s["flags"] & 16 for s in spans)
                        style = 'Heading 1' if (max_size > 13 or is_bold) else 'Normal'
                        paragraphs.append({
                            'text': text,
                            'style': style,
                            'bold': is_bold,
                            'font_size': max_size
                        })
                elif block["type"] == 1 and soa_image_url is None:
                    prev_text_block = False
                    try:
                        xref = block.get("xref")
                        if xref:
                            img = doc.extract_image(xref)
                            soa_image_url = self._save_image(img["image"], img.get("ext", "png"), upload_dir)
                    except Exception:
                        pass
                else:
                    prev_text_block = False

        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    for tbl in (page.extract_tables() or []):
                        cleaned = [[str(c or '').strip() for c in row]
                                   for row in tbl if any(c for c in row)]
                        if len(cleaned) >= 2:
                            tables_data.append(cleaned)
        except Exception as e:
            logger.warning(f"pdfplumber: {e}")

        # Extract raw TOC text for LLM agent (using pypdf outline + page scan)
        raw_toc_text = ''
        try:
            reader = PdfReader(io.BytesIO(content))
            raw_toc_text = self._extract_raw_toc_text(reader)
        except Exception as e:
            logger.warning(f"TOC raw extraction failed: {e}")

        full_text = '\n'.join(p['text'] for p in paragraphs)
        return self._run(paragraphs, full_text, tables_data, soa_image_url, raw_toc_text=raw_toc_text)

    # ──────────────────────────────────────────────
    # MAIN PIPELINE
    # ──────────────────────────────────────────────
    def _run(self, paragraphs, full_text, tables_data, soa_image_url, raw_toc_text: str = '') -> dict:
        # STEP 1: Extract all sections + subsections structurally (Python, no AI)
        sections = self._extract_sections_structural(paragraphs, full_text)

        if GROQ_API_KEY:
            try:
                result = self._groq_pipeline(paragraphs, full_text, tables_data,
                                              soa_image_url, sections, raw_toc_text)
                logger.info("Groq extraction succeeded")
                return result
            except Exception as e:
                logger.warning(f"Groq failed, using regex fallback: {e}")

        logger.info("Using regex fallback")
        return self._regex_all(paragraphs, full_text, tables_data, soa_image_url, sections)

    @staticmethod
    def _subsection_heading_from_line(text: str) -> Optional[Tuple[str, str, str]]:
        """
        Detect a numbered subsection at the start of a line.
        Returns (number_str, title, overflow) where overflow is prose merged on the same PDF line.
        """
        m = SUBSECTION_HEAD_RE.match(text.strip())
        if not m:
            return None
        num_str, rest = m.group(1), m.group(2).strip()
        if re.search(r'\d+\.\d+\s*(?:mg|ml|µg|mcg|g)\b', text, re.I):
            return None
        first_w = (rest.split(None, 1)[0].lower() if rest else '')
        if first_w in ('million', 'billion', 'thousand', 'percent', 'fold', 'times'):
            return None
        overflow = ''
        if len(rest) > 180:
            cut = rest[:360]
            boundary = None
            for sep in ('. The ', '. This ', '. These ', '. An ', '. A ', '. If ', '. All '):
                pos = cut.find(sep)
                if pos > 35:
                    boundary = pos + 1
                    break
            if boundary is None:
                bm = re.search(r'(?<=[a-z\)])\. (?=[A-Z(])', cut)
                if bm and bm.start() > 35:
                    boundary = bm.start() + 1
            if boundary:
                overflow = rest[boundary:].lstrip('. ')
                rest = rest[:boundary].strip().rstrip('.')
        return num_str, rest, overflow

    # ──────────────────────────────────────────────
    # STEP 1: STRUCTURAL SECTION + SUBSECTION EXTRACTION
    # Handles TOC skipping, any number of sections (up to 30), subsections
    # ──────────────────────────────────────────────
    def _extract_sections_structural(self, paragraphs: List[Dict], full_text: str) -> dict:
        """
        3-phase section extraction:
          Phase 1 – Find and skip the Table of Contents page
          Phase 2 – Scan for all headings, classify each as top-level / subsection
          Phase 3 – Build content under each heading, flush into sections dict

        Subsections use {title, content, customTable} matching ProtocolSections.jsx.
        SoA / Trial Procedures / Assessments headings are always captured as sections.
        """

        # ─── TOC detection patterns ───
        TOC_DOTS_RE  = re.compile(r'[.\-–]{3,}\s*\d+\s*$')      # "Introduction ......... 3"
        TOC_TITLE_KW = {'table of contents', 'contents', 'toc', 'table of content'}
        # Keywords that MUST always become a section regardless of location
        SOA_KW = {'schedule of activities', 'trial procedures', 'assessments and procedures',
                  'study procedures', 'schedule of events', 'soa', 'study activities',
                  'time and events schedule', 'table 1 time and events schedule'}

        # ─────────────────────────────────────────────
        # Phase 1: Identify the TOC region and skip it.
        # The TOC is the block of text between a "TABLE OF CONTENTS" heading and
        # the first heading that has substantial prose content after it.
        # ─────────────────────────────────────────────
        toc_region: set = set()   # set of paragraph indices that are part of the TOC
        toc_start = -1
        toc_end   = -1

        for i, p in enumerate(paragraphs):
            t_lower = p['text'].strip().lower()
            # Detect explicit TOC heading
            if t_lower in TOC_TITLE_KW or t_lower.startswith('table of content'):
                toc_start = i
                break

        if toc_start >= 0:
            # Walk forward from TOC heading:
            # The TOC ends at the first paragraph that looks like a real section heading
            # that is followed by ≥1 lines of prose (not just more headings / TOC entries)
            for i in range(toc_start + 1, len(paragraphs)):
                p = paragraphs[i]
                text = p['text'].strip()
                # Is it a numbered / keyword heading?
                is_num_head = bool(SECTION_NUMBER_RE.match(text) or SUBSECTION_RE.match(text))
                is_kw_head  = ('Heading' in p.get('style','') or p.get('bold', False) and len(text) < 150)

                # Is it a TOC-style line (dots + page number)?
                is_toc_entry = bool(TOC_DOTS_RE.search(text))
                if is_toc_entry:
                    toc_region.add(i)
                    continue

                if is_num_head or is_kw_head:
                    # Peek ahead: count real-prose lines in next 6 paragraphs
                    prose_count = 0
                    for j in range(i + 1, min(i + 7, len(paragraphs))):
                        nxt = paragraphs[j]['text'].strip()
                        if (not SECTION_NUMBER_RE.match(nxt)
                                and not SUBSECTION_RE.match(nxt)
                                and not TOC_DOTS_RE.search(nxt)
                                and len(nxt) > 40):
                            prose_count += 1
                    if prose_count >= 1:
                        toc_end = i
                        break
                    else:
                        toc_region.add(i)  # heading with no prose = TOC entry

        # All indices from toc_start to toc_end (exclusive) are TOC
        if toc_start >= 0 and toc_end > toc_start:
            for i in range(toc_start, toc_end):
                toc_region.add(i)

        # ─────────────────────────────────────────────
        # Phase 2 + 3: Parse real document content
        # ─────────────────────────────────────────────
        sections: Dict[str, Dict] = {}

        # If a TOC was found, create it as the very first special section
        if toc_region:
            toc_lines = []
            for i in sorted(list(toc_region)):
                t_text = paragraphs[i]['text'].strip()
                if t_text:
                    toc_lines.append(t_text)
            if toc_lines:
                sections['0'] = {
                    'title': 'Table of Contents',
                    'main': '<br/>\n'.join(toc_lines),
                    'notes': '',
                    'subsections': []
                }

        current_top_id: Optional[str] = None
        current_top_title: str = ''
        current_sub_title: str = ''
        current_sub_content: List[str] = []
        current_top_content: List[str] = []
        is_in_sub: bool = False
        # Track all encountered section IDs so we can auto-assign unknown headings
        next_auto_id: List[int] = [100]   # auto IDs start at 100 to avoid collision

        def _format_prose(text: str) -> str:
            lines = text.split('\n')
            res = []
            for i, line in enumerate(lines):
                line = line.strip()
                if not line: continue
                is_bullet = re.match(r'^([\u2022\u00b7\-\*]|\d+[\.\)])\s+', line)
                if is_bullet:
                    if len(res) > 0:
                        res.append('<br/>\n' + html_module.escape(line))
                    else:
                        res.append(html_module.escape(line))
                else:
                    if len(res) > 0 and not res[-1].endswith('<br/>\n'):
                        res[-1] += ' ' + html_module.escape(line)
                    else:
                        res.append(html_module.escape(line))
            return ''.join(res)

        def _flush_sub():
            nonlocal current_sub_content, is_in_sub, current_sub_title
            if current_top_id and current_sub_title:
                body = '\n'.join(current_sub_content).strip()
                if body:
                    paras = [x.strip() for x in re.split(r'\n\s*\n+', body) if x.strip()]
                    frag = ''.join(f'<p>{_format_prose(p)}</p>' for p in paras)
                    html_body = f"<h3>{html_module.escape(current_sub_title)}</h3>\n{frag}"
                else:
                    html_body = f"<h3>{html_module.escape(current_sub_title)}</h3>"
                sections[current_top_id]['subsections'].append({
                    'title': current_sub_title,
                    'content': html_body,
                    'customTable': None
                })
            current_sub_content.clear()
            current_sub_title = ''
            is_in_sub = False

        def _flush_top():
            nonlocal current_top_content
            if current_top_id and current_top_content:
                body = '\n'.join(current_top_content).strip()
                paras = [x.strip() for x in re.split(r'\n\s*\n+', body) if x.strip()]
                body_html = ''.join(f'<p>{_format_prose(p)}</p>' for p in paras)
                existing = sections.get(current_top_id, {}).get('main', '')
                merged = (existing + '\n' + body_html).strip() if body_html else existing
                sections[current_top_id]['main'] = merged
            current_top_content.clear()

        def _new_section(sid: str, title: str):
            nonlocal current_top_id, current_top_title, is_in_sub
            _flush_sub()
            _flush_top()
            current_top_id    = sid
            current_top_title = title
            is_in_sub = False
            if sid not in sections:
                sections[sid] = {'title': title, 'main': '', 'notes': '', 'subsections': []}

        for i, p in enumerate(paragraphs):
            # Skip TOC region
            if i in toc_region:
                continue

            text  = _strip_template_noise(p['text'].strip())
            style = p.get('style', 'Normal')
            bold  = p.get('bold', False)
            fsize = p.get('font_size') or 12

            if not text:
                if current_top_id:
                    if is_in_sub:
                        current_sub_content.append('')
                    else:
                        current_top_content.append('')
                continue

            # Skip standalone page numbers
            if re.match(r'^\d{1,3}$', text):
                continue

            # ─── Check for subsection heading (e.g. "1.1 Background", "8.3.1 Definition …") ───
            sub_parse = self._subsection_heading_from_line(text)
            if sub_parse:
                num_str, sub_head, same_line_overflow = sub_parse
                top_num = num_str.split('.')[0]

                if top_num in sections:
                    if current_top_id != top_num:
                        _flush_sub()
                        _flush_top()
                        current_top_id = top_num
                        is_in_sub = False
                    _flush_sub()
                    current_sub_title = sub_head
                    is_in_sub = True
                    if same_line_overflow:
                        current_sub_content.append(same_line_overflow)
                    continue
                elif top_num not in sections:
                    _new_section(top_num, f'Section {top_num}')
                    _flush_sub()
                    current_sub_title = sub_head
                    is_in_sub = True
                    if same_line_overflow:
                        current_sub_content.append(same_line_overflow)
                    continue

            # ─── Check for numbered top-level heading (e.g. "2. Introduction") ───
            num_m = SECTION_NUMBER_RE.match(text)
            matched_id    = None
            matched_title = text

            if num_m:
                num_val = int(num_m.group(1))
                if 1 <= num_val <= 30:
                    cand_title = num_m.group(2).strip()
                    # A true numbered section heading must be styled, bold, uppercase, or relatively short
                    # ALSO: reject if the candidate title looks like a date, exclusion criterion or sentence
                    _is_garbage_title = bool(
                        re.match(r'^(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\b', cand_title, re.I)
                        or re.match(r'^\d{1,2}\s+(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)', cand_title, re.I)
                        or (len(cand_title) > 80 and not bold and 'Heading' not in style)
                        or re.match(r'^(Any|All|No |Patients|Subjects|Smokers|History|Prior|Use of|Current)', cand_title, re.I)
                    )
                    if not _is_garbage_title and (len(text) < 80 or bold or 'Heading' in style or text.isupper()):
                        matched_id    = str(num_val)
                        matched_title = cand_title

            # ─── Check for keyword / style-based heading ───
            if not matched_id:
                is_style_head = (
                    'Heading' in style
                    or (bold and len(text) < 150 and fsize >= 11)
                    or (text.isupper() and 5 < len(text) < 150)
                )
                if is_style_head:
                    text_upper = text.upper()
                    # SoA keywords always get their own section
                    if any(kw in text_upper.lower() for kw in SOA_KW):
                        matched_id    = str(next_auto_id[0])
                        next_auto_id[0] += 1
                        matched_title = text
                    else:
                        for keyword, sid in SECTION_HEADING_MAP.items():
                            if keyword in text_upper:
                                matched_id    = sid
                                matched_title = text
                                break

            # ─── Open the new section ───
            if matched_id:
                _new_section(matched_id, matched_title)
                continue

            # ─── Regular content line ───
            if is_in_sub:
                current_sub_content.append(text)
            elif current_top_id:
                current_top_content.append(text)

        # Final flush
        _flush_sub()
        _flush_top()

        # Post-process: remove any section stubs (no content AND no subsections)
        # typically these are leftover TOC-like entries that slipped through
        to_remove = [
            sid for sid, sec in sections.items()
            if not sec.get('main', '').strip() and not sec.get('subsections')
        ]
        for sid in to_remove:
            del sections[sid]

        return sections



    # ──────────────────────────────────────────────
    # GROQ PIPELINE: 4 focused calls
    # ──────────────────────────────────────────────
    def _groq_pipeline(self, paragraphs, full_text, tables_data, soa_image_url, sections, raw_toc_text: str = '') -> dict:
        from groq import Groq
        client = Groq(api_key=GROQ_API_KEY)

        # ── Prepare text slices ──
        # Reduce character budgets drastically (approx. 4 chars per token)
        cover_text = self._smart_cover(paragraphs, full_text, budget=4000)
        synopsis_text = self._smart_synopsis_text(full_text, budget=6000)
        table_text = self._tables_to_pipe(tables_data, budget=8000)

        # ── Call 1: Metadata ──
        meta = self._groq_call(client, META_SYSTEM, META_PROMPT.format(text=cover_text),
                               max_tokens=800, label="Metadata")

        # ── Call 2: Synopsis / criteria ──
        syn = self._groq_call(client, SYN_SYSTEM, SYN_PROMPT.format(text=synopsis_text),
                              max_tokens=1000, label="Synopsis")

        # ── Call 3: SoA ──
        soa = {"headers": [], "rows": []}
        if table_text.strip():
            soa = self._groq_call(client, SOA_SYSTEM, SOA_PROMPT.format(tables=table_text),
                                  max_tokens=3000, label="SoA")

        # ── Call 4: TOC (LLM agent — array response, not json_object mode) ──
        toc_tree = []
        if raw_toc_text.strip():
            toc_tree = self._groq_toc_call(client, raw_toc_text)

        return self._assemble(meta, syn, soa, sections, soa_image_url, full_text, toc_tree, tables_data=tables_data)

    def _groq_call(self, client, system: str, user: str, max_tokens: int, label: str) -> dict:
        """Single Groq call with JSON mode. Returns parsed dict."""
        try:
            resp = client.chat.completions.create(
                model=GROQ_MODEL,
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user",   "content": user},
                ],
                temperature=0.05,
                max_tokens=max_tokens,
                response_format={"type": "json_object"},
            )
            raw = resp.choices[0].message.content.strip()
            raw = re.sub(r'^```(?:json)?\s*', '', raw, flags=re.MULTILINE)
            raw = re.sub(r'```\s*$', '', raw, flags=re.MULTILINE)
            return json.loads(raw.strip())
        except Exception as e:
            logger.warning(f"Groq {label} call failed: {e}")
            return {}

    def _groq_toc_call(self, client, raw_toc_text: str) -> list:
        """
        Dedicated Groq call for Table of Contents.
        Returns a list (array) — cannot use json_object mode since the root is [].
        """
        try:
            prepared = self._prepare_raw_toc_for_llm(raw_toc_text)
            user_prompt = TOC_USER_TEMPLATE.format(raw_toc=prepared)
            resp = client.chat.completions.create(
                model=GROQ_MODEL,
                messages=[
                    {"role": "system", "content": TOC_SYSTEM},
                    {"role": "user",   "content": user_prompt},
                ],
                temperature=0.0,
                max_tokens=2000,
            )
            raw = resp.choices[0].message.content.strip()
            # Strip accidental markdown fences
            if raw.startswith('```'):
                parts = raw.split('```')
                raw = parts[1] if len(parts) > 1 else raw
                if raw.startswith('json'):
                    raw = raw[4:]
            raw = raw.strip()
            result = self._parse_toc_json_array(raw)
            return result if isinstance(result, list) else []
        except Exception as e:
            logger.warning(f"Groq TOC call failed: {e}")
            return []

    @staticmethod
    def _parse_toc_json_array(raw: str) -> list:
        """Parse LLM TOC response; tolerate stray text around a JSON array."""
        raw = raw.strip()
        try:
            result = json.loads(raw)
            return result if isinstance(result, list) else []
        except json.JSONDecodeError:
            pass
        start = raw.find('[')
        end = raw.rfind(']')
        if start >= 0 and end > start:
            try:
                result = json.loads(raw[start : end + 1])
                return result if isinstance(result, list) else []
            except json.JSONDecodeError:
                logger.warning("TOC JSON array slice parse failed")
        return []

    @staticmethod
    def _prepare_raw_toc_for_llm(raw_toc_text: str, budget: int = 56000) -> str:
        """
        Keep the most useful TOC bytes within the model context limit.
        Printed TOC pages (dot leaders) are kept at the start of the prompt.
        """
        s = (raw_toc_text or '').strip()
        if len(s) <= budget:
            return s
        marker_pages = '--- TOC PAGES ---'
        marker_bm = '--- PDF BOOKMARKS ---'
        if marker_pages in s:
            idx = s.find(marker_pages)
            prefix = s[: idx + len(marker_pages)]
            rest = s[idx + len(marker_pages) :].lstrip('\n')
            if len(prefix) >= budget:
                return prefix[:budget]
            take = budget - len(prefix) - 2
            return (prefix + '\n\n' + rest[:take]).strip()[:budget]
        if marker_bm in s:
            idx = s.find(marker_bm)
            after = s[idx:]
            return after[:budget]
        return s[:budget]

    @staticmethod
    def _count_toc_dot_leader_lines(text: str) -> int:
        if not text:
            return 0
        return sum(
            1
            for line in text.splitlines()
            if re.search(r'[.\-–]{3,}\s*\d+\s*$', line.strip())
        )

    @staticmethod
    def _scanned_toc_body(scanned_text: str) -> str:
        if not scanned_text:
            return ''
        return scanned_text.split('--- TOC PAGES ---', 1)[-1].strip()

    @staticmethod
    def _scanned_toc_looks_valid(scanned_text: str) -> bool:
        """True if extracted pages look like a real TOC (not a random body page)."""
        body = DocumentParser._scanned_toc_body(scanned_text)
        if not body:
            return False
        if DocumentParser._count_toc_dot_leader_lines(body) >= 2:
            return True
        numbered = sum(
            1
            for line in body.splitlines()
            if re.match(r'^\s*\d+(?:\.\d+)*[\.\)\s]', line.strip())
        )
        return numbered >= 5

    # ──────────────────────────────────────────────
    # SMART TEXT SLICES
    # ──────────────────────────────────────────────
    def _smart_cover(self, paragraphs, full_text, budget=15000) -> str:
        """First N chars — almost always contains title, number, PI, sponsor, date, version."""
        return full_text[:budget]

    def _smart_synopsis_text(self, full_text: str, budget=12000) -> str:
        """
        Slice text around key synopsis keywords so Groq sees only the relevant parts.
        This keeps token count low regardless of document size.
        """
        targets = [
            ('objectives and endpoints', 3500),
            ('primary objective', 900),
            ('secondary objective', 800),
            ('exploratory objective', 600),
            ('primary endpoint', 900),
            ('secondary endpoint', 800),
            ('co-primary endpoint', 800),
            ('key secondary endpoint', 800),
            ('main secondary endpoint', 800),
            ('efficacy endpoint', 800),
            ('inclusion criteria', 900),
            ('exclusion criteria', 900),
            ('statistical method', 400),
            ('sample size', 300),
            ('number of patient', 300),
            ('exploratory endpoint', 600),
        ]
        parts = []
        used = 0
        text_lower = full_text.lower()

        # Prefer a full Section 3 block when headings are visible (ICH-style protocols)
        sec3 = re.search(
            r'(?is)(?:^|\n)\s*3[\.\)\s]+.{0,120}?(?:OBJECTIVES|ENDPOINTS).+?(?=\n\s*4[\.\)\s]+)',
            full_text[:250000],
        )
        if sec3:
            chunk = sec3.group(0).strip()
            if len(chunk) > 400:
                tag = f"\n[SECTION_3_OBJECTIVES_ENDPOINTS]\n{chunk[:8000]}"
                parts.append(tag)
                used += len(tag)

        for keyword, allot in targets:
            if used >= budget:
                break
            idx = text_lower.find(keyword)
            if idx == -1:
                continue
            snippet = full_text[max(0, idx - 120): idx + allot]
            tag = f"\n[{keyword.upper()}]\n{snippet}"
            parts.append(tag)
            used += len(tag)

        return '\n'.join(parts)[:budget] if parts else full_text[:budget]

    def _tables_to_pipe(self, tables_data: List, budget: int) -> str:
        """Format all extracted tables as pipe-separated text for Groq."""
        parts = []
        for i, tbl in enumerate(tables_data[:20]):
            rows_text = [" | ".join(str(c) for c in row) for row in tbl]
            parts.append(f"TABLE {i+1}:\n" + "\n".join(rows_text))
        return "\n\n".join(parts)[:budget]

    # ──────────────────────────────────────────────
    # TOC RAW TEXT EXTRACTION (pypdf)
    # ──────────────────────────────────────────────
    def _extract_raw_toc_text(self, reader) -> str:
        """
        Pull the best available TOC representation from the PDF.
        Priority:
          A) PDF outline/bookmark tree   — flatten to indented text
          B) TOC page(s) raw text        — send directly to Groq
        """
        outline = reader.outline

        # ── A: Outline exists ──────────────────────────────────────────────
        if outline and len(outline) > 0:
            lines = []
            self._flatten_outline(outline, reader, lines, depth=0)
            return "\n".join(lines)

        # ── B: No outline — find TOC page(s) by scanning text ──────────────
        toc_chunks = []
        in_toc = False

        for i, page in enumerate(reader.pages[:50]):
            text = (page.extract_text() or "").strip()
            if not text:
                continue

            lower = text.lower()

            # Detect TOC page start
            if not in_toc and ("table of contents" in lower or "contents" in lower):
                in_toc = True

            if in_toc:
                toc_chunks.append(text)
                # Stop after 10 pages once inside TOC (multi-page TOCs)
                if len(toc_chunks) >= 10:
                    break
                # Stop if we see a section heading that clearly is no longer TOC
                if len(toc_chunks) > 1 and self._looks_like_body_page(text):
                    toc_chunks.pop()  # remove the body page we just added
                    break

        if toc_chunks:
            return "\n\n--- PAGE BREAK ---\n\n".join(toc_chunks)

        # Fallback: first 10 pages
        fallback = []
        for page in reader.pages[:10]:
            fallback.append(page.extract_text() or "")
        return "\n".join(fallback)

    def _flatten_outline(self, items, reader, lines, depth):
        """Recursively flatten PDF outline into indented lines for Groq."""
        for item in items:
            if isinstance(item, list):
                self._flatten_outline(item, reader, lines, depth + 1)
            else:
                title = item.get('/Title', '').strip()
                try:
                    page_num = reader.get_destination_page_number(item) + 1
                    lines.append(f"{'  ' * depth}{title}  [page {page_num}]")
                except Exception:
                    if title:
                        lines.append(f"{'  ' * depth}{title}")

    def _looks_like_body_page(self, text: str) -> bool:
        """Heuristic: is this page actual document body rather than TOC?"""
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        # Body pages typically have long prose sentences
        long_lines = sum(1 for l in lines if len(l) > 120)
        return long_lines >= 3

    # ──────────────────────────────────────────────
    # ASSEMBLE FINAL RESULT
    # ──────────────────────────────────────────────
    def _assemble(self, meta, syn, soa, sections, soa_image_url, full_text, toc_tree: list = None, tables_data: list = None) -> dict:
        self._last_tables_data = tables_data or []
        result = self._empty_protocol()

        s   = lambda d, k: str(d.get(k) or '').strip()
        lst = lambda d, k: [str(v).strip() for v in (d.get(k) or []) if v and str(v).strip()]

        # ── Scalars from meta ──
        result['protocol_title']          = s(meta, 'protocol_title')
        result['protocol_number']         = s(meta, 'protocol_number')
        result['nct_number']              = s(meta, 'nct_number')
        result['principal_investigator']  = s(meta, 'principal_investigator')
        result['sponsor']                 = s(meta, 'sponsor')
        result['funded_by']               = s(meta, 'funded_by')
        result['version_number']          = s(meta, 'version_number') or 'v1.0'
        result['protocol_date']           = s(meta, 'protocol_date')

        # ── Approval details ──
        d = result['approval_data']['details']
        d['protocol_name']              = result['protocol_title']
        d['protocol_number']            = result['protocol_number']
        d['imp']                        = s(meta, 'imp')
        d['indication']                 = s(meta, 'indication')
        d['clinical_phase']             = s(meta, 'clinical_phase')
        d['investigators']              = result['principal_investigator']
        d['coordinating_investigator']  = s(meta, 'coordinating_investigator')
        d['expert_committee']           = s(meta, 'expert_committee')
        d['sponsor_name_address']       = s(meta, 'sponsor_name_address')

        # ── Synopsis overview ──
        ov = result['synopsis_data']['overview']
        ov['title']                     = result['protocol_title']
        ov['coordinating_investigator'] = s(meta, 'coordinating_investigator')
        ov['expert_committee']          = s(meta, 'expert_committee')
        ov['investigators']             = result['principal_investigator']
        ov['trial_sites']               = s(meta, 'trial_sites')
        ov['planned_period']            = s(meta, 'planned_period')
        ov['fpfv']                      = s(meta, 'fpfv')
        ov['lplv']                      = s(meta, 'lplv')
        ov['clinical_phase']            = s(meta, 'clinical_phase')

        # ── Objectives / Endpoints from synopsis call ──
        result['synopsis_data']['objectives']['primary']     = lst(syn, 'primary_objectives')
        result['synopsis_data']['objectives']['secondary']   = lst(syn, 'secondary_objectives')
        result['synopsis_data']['objectives']['exploratory'] = lst(syn, 'exploratory_objectives')
        result['synopsis_data']['endpoints']['primary']      = lst(syn, 'primary_endpoints')
        result['synopsis_data']['endpoints']['secondary']    = lst(syn, 'secondary_endpoints')
        result['synopsis_data']['endpoints']['exploratory']  = lst(syn, 'exploratory_endpoints')
        result['synopsis_data']['num_patients']              = s(meta, 'num_patients')
        result['synopsis_data']['statistical_methods']       = s(syn, 'statistical_methods')

        incl = lst(syn, 'inclusion_criteria')
        excl = lst(syn, 'exclusion_criteria')
        result['synopsis_data']['inclusion'] = {'text': '\n'.join(incl), 'points': incl}
        result['synopsis_data']['exclusion'] = {'text': '\n'.join(excl), 'points': excl}

        # ── Populate Section 3 (Objectives/Endpoints) Table ──
        prim_objs = lst(syn, 'primary_objectives')
        prim_ends = lst(syn, 'primary_endpoints')
        sec_objs = lst(syn, 'secondary_objectives')
        sec_ends = lst(syn, 'secondary_endpoints')
        exp_objs = lst(syn, 'exploratory_objectives')
        exp_ends = lst(syn, 'exploratory_endpoints')

        s3_rows = []
        if prim_objs or prim_ends:
            s3_rows.append(['Primary', '<br/>'.join(f"• {o}" for o in prim_objs), '<br/>'.join(f"• {e}" for e in prim_ends)])
        if sec_objs or sec_ends:
            s3_rows.append(['Secondary', '<br/>'.join(f"• {o}" for o in sec_objs), '<br/>'.join(f"• {e}" for e in sec_ends)])
        if exp_objs or exp_ends:
            s3_rows.append(['Exploratory', '<br/>'.join(f"• {o}" for o in exp_objs), '<br/>'.join(f"• {e}" for e in exp_ends)])
        
        result['section3']['table']['rows'] = s3_rows

        # ── Populate synopsis_data UI fields (matches frontend ProtocolContext key) ──
        # Map extracted objectives/endpoints into the synopsis_data structure the UI reads
        result['synopsis_data'] = result.get('synopsis_data', {})
        existing = result['synopsis_data']
        # Populate objectives (nested by type)
        if prim_objs or sec_objs or exp_objs:
            existing.setdefault('objectives', {})
            if prim_objs: existing['objectives']['primary'] = prim_objs
            if sec_objs:  existing['objectives']['secondary'] = sec_objs
            if exp_objs:  existing['objectives']['exploratory'] = exp_objs
        # Populate endpoints
        if prim_ends or sec_ends or exp_ends:
            existing.setdefault('endpoints', {})
            if prim_ends: existing['endpoints']['primary'] = prim_ends
            if sec_ends:  existing['endpoints']['secondary'] = sec_ends
            if exp_ends:  existing['endpoints']['exploratory'] = exp_ends
        # Populate overview fields from meta
        existing.setdefault('overview', {})
        if not existing['overview'].get('title'):   existing['overview']['title']   = result.get('protocol_title', '')
        if not existing['overview'].get('clinical_phase'): existing['overview']['clinical_phase'] = s(meta, 'clinical_phase')
        if not existing['overview'].get('trial_sites'):    existing['overview']['trial_sites']    = s(meta, 'trial_sites')
        if not existing['overview'].get('planned_period'): existing['overview']['planned_period']  = s(meta, 'planned_period')
        if not existing.get('num_patients'):               existing['num_patients']               = s(meta, 'num_patients')
        if not existing.get('imp'):                        existing['imp']                        = s(meta, 'imp')

        # ── Sections (structural, unlimited) ──
        result['sections'] = sections

        # ── TOC tree from LLM agent (if available) ──
        if toc_tree:
            if '0' not in result['sections']:
                result['sections']['0'] = {
                    'title': 'Table of Contents',
                    'main': '',
                    'notes': '',
                    'subsections': []
                }
            result['sections']['0']['toc_tree'] = toc_tree

            # Sync TOC tree into Protocol Sections (update titles + create stubs)
            self._sync_sections_from_toc(result['sections'], toc_tree)

            # Collect ALL valid section numbers from the entire TOC tree (recursive)
            valid_keys = {'0'}  # Always keep the TOC section itself
            self._collect_toc_section_nums(toc_tree, valid_keys)

            # Aggressively purge every section key that the TOC didn't authorize.
            # This removes both numeric garbage (e.g. "23") AND string-key garbage
            # (e.g. 'synopsis', 'withdrawal', 'supporting') that slipped in from the
            # structural heuristic parser.
            ALWAYS_KEEP = {'0'}  # key '0' = TOC section
            to_remove = [
                k for k in list(result['sections'].keys())
                if k not in ALWAYS_KEEP and k not in valid_keys
            ]
            for k in to_remove:
                del result['sections'][k]

        # ── SoA table ──
        CHECKS = {'1','x','✓','✔','y','yes','true','•','×','xi','checked','v'}
        g_hdrs = soa.get('headers', [])
        g_rows = soa.get('rows', [])

        if isinstance(g_hdrs, list) and g_hdrs and isinstance(g_rows, list) and g_rows:
            clean_hdrs = [str(h).strip() for h in g_hdrs if str(h).strip()]
            if clean_hdrs and clean_hdrs[0].lower() not in ['procedure', 'procedures', 'assessment', 'assessments']:
                clean_hdrs.insert(0, 'Procedure')
                
            clean_rows = []
            for row in g_rows:
                if not isinstance(row, list) or not row:
                    continue
                proc = str(row[0]).strip()
                expected = max(0, len(clean_hdrs) - 1)
                cells = []
                for c in row[1:]:
                    cv = str(c).strip().lower()
                    cells.append('1' if cv in CHECKS else '0')
                while len(cells) < expected:
                    cells.append('0')
                clean_rows.append([proc] + cells[:expected])
            if clean_rows:
                result['soa_data']['table']['headers'] = clean_hdrs
                result['soa_data']['table']['rows']    = clean_rows
        else:
            # Regex SoA fallback — use the raw tables_data from pdfplumber
            result['soa_data'] = self._regex_soa(
                self._last_tables_data or self._all_tables_as_data(full_text), soa_image_url
            )

        # Attach image if table empty
        if not result['soa_data']['table']['rows'] and soa_image_url:
            result['soa_data']['image'] = {
                'url': soa_image_url,
                'caption': 'Schedule of Activities (extracted from document)',
                'description': ''
            }

        return result

    # ──────────────────────────────────────────────
    # SYNC TOC TREE → PROTOCOL SECTIONS
    # ──────────────────────────────────────────────
    def _sync_sections_from_toc(self, sections: dict, toc_tree: list) -> None:
        """
        Walk the LLM-parsed TOC tree and upsert into the `sections` dict so that
        Protocol Sections reflects the clean TOC hierarchy:

        Rules:
          • Only numbered sections (number != "") are used (unnumbered like REFERENCES
            at the top-level are kept as-is from structural extraction).
          • Top-level items (level 0, single integer number "1", "2", …) map to
            sections dict keys "1", "2", …
          • If a section already exists, update its title with the clean TOC title.
          • If a section does NOT exist, create an empty stub so it appears in the grid.
          • Children (level 1+) become subsections of their parent section.
          • Existing subsections are updated by index match; extra TOC subsections
            are appended.
        """
        # Collect auto-ID counter for unnumbered top-level stubs (skip section 0)
        next_auto = max((int(k) for k in sections if k.lstrip('-').isdigit() and int(k) > 0),
                        default=0) + 1

        for node in toc_tree:
            num = (node.get('number') or '').strip()
            title = (node.get('title') or '').strip()
            children = node.get('children') or []

            if not title:
                continue

            # ── Determine section key ──
            if num and re.match(r'^\d+$', num):          # "1", "2", … → direct key
                sid = num
            elif num and re.match(r'^\d+\.', num):       # "1.1" → belongs to parent '1', not a top-level
                # Subsection at root level in toc_tree (shouldn’t happen but handle gracefully)
                parent_num = num.split('.')[0]
                if parent_num in sections:
                    self._upsert_subsection(sections[parent_num], num, title, children)
                continue
            else:
                # Unnumbered top-level (e.g. APPROVAL, SYNOPSIS, REFERENCES)
                # Try to map via SECTION_HEADING_MAP so they get a real section key
                title_upper = title.upper()
                mapped_sid = None
                for kw, map_sid in SECTION_HEADING_MAP.items():
                    if kw in title_upper:
                        mapped_sid = map_sid
                        break
                if mapped_sid:
                    sid = mapped_sid
                else:
                    # Truly unknown unnumbered section — structural parser handles it
                    continue

            # ── Upsert the top-level section ──
            if sid not in sections:
                # Create stub — structural parser missed this section entirely
                sections[sid] = {
                    'title': title,
                    'main': '',
                    'notes': '',
                    'subsections': []
                }
            else:
                # Update title to clean TOC version (overrides heuristic title)
                sections[sid]['title'] = title
                if 'subsections' not in sections[sid]:
                    sections[sid]['subsections'] = []

            # ── Sync children into subsections ──
            if children:
                self._sync_subsections_from_toc(sections[sid], children)

    def _sync_subsections_from_toc(self, section: dict, children: list) -> None:
        """
        Merge TOC children into a section’s flat subsections list in exact TOC order.
        To maintain UI compatibility while supporting deep nesting, we flatten
        grandchildren into the section's array but assign them `depth` and `number`.
        """
        existing_subs = section.setdefault('subsections', [])

        def build_ordered(nodes, depth=1):
            ordered = []
            for child in nodes:
                num   = (child.get('number') or '').strip()
                title = (child.get('title') or '').strip()
                grandchildren = child.get('children') or []

                if not title:
                    continue

                # Try to find existing subsection by title match (case-insensitive)
                matched_idx = None
                title_lower = title.lower()
                for i, sub in enumerate(existing_subs):
                    if sub is None: 
                        continue
                    existing_lower = (sub.get('title') or '').lower()
                    # Exact or prefix match
                    if existing_lower == title_lower or existing_lower.startswith(title_lower[:20]):
                        matched_idx = i
                        break

                if matched_idx is not None:
                    # Update title to clean TOC version
                    matched_sub = existing_subs[matched_idx]
                    existing_subs[matched_idx] = None  # mark as consumed
                    matched_sub['title'] = title
                    matched_sub['number'] = num
                    matched_sub['depth'] = depth
                    ordered.append(matched_sub)
                else:
                    # Append new empty stub subsection
                    ordered.append({
                        'title': title,
                        'content': '',
                        'customTable': None,
                        'number': num,
                        'depth': depth
                    })

                # Recurse for grandchildren, extending the flat ordered list
                if grandchildren:
                    ordered.extend(build_ordered(grandchildren, depth + 1))
            
            return ordered

        # Construct the perfectly ordered list from the TOC tree
        ordered_subs = build_ordered(children, 1)

        # Merge the text content of any heuristic "leftovers" into the main section body
        # NO: we keep them because TOC misses legitimate sub-sections sometimes.
        for sub in existing_subs:
            if sub is not None:
                ordered_subs.append(sub)

        section['subsections'] = ordered_subs

    def _upsert_subsection(self, section: dict, num: str, title: str, children: list) -> None:
        """Helper to upsert a single subsection by number into a section."""
        existing_subs = section.setdefault('subsections', [])
        title_lower = title.lower()
        for sub in existing_subs:
            if (sub.get('title') or '').lower() == title_lower:
                sub['title'] = title
                return
        existing_subs.append({'title': title, 'content': '', 'customTable': None})

    def _collect_toc_section_nums(self, nodes: list, out: set) -> None:
        """
        Recursively walk the LLM TOC tree and collect every top-level integer section
        number (e.g. "1", "2", ... "17") into `out`.
        Also maps well-known unnumbered front-matter items (SYNOPSIS, APPROVAL, etc.)
        to their corresponding heuristic section IDs via SECTION_HEADING_MAP.
        """
        for node in nodes:
            num = (node.get('number') or '').strip()
            title_upper = (node.get('title') or '').upper()

            # Collect top-level integer keys
            if num and re.match(r'^\d+$', num):
                out.add(num)

            # Map unnumbered front-matter headings to their heuristic section IDs
            for kw, sid in SECTION_HEADING_MAP.items():
                if kw in title_upper:
                    out.add(sid)
                    break

            # Recurse into children (to catch any nested numbered sections too)
            children = node.get('children') or []
            if children:
                self._collect_toc_section_nums(children, out)


    # ──────────────────────────────────────────────
    # REGEX FALLBACK (full)
    # ──────────────────────────────────────────────
    def _regex_all(self, paragraphs, full_text, tables_data, soa_image_url, sections) -> dict:
        result = self._empty_protocol()

        phase = self._rx(full_text, [r'Phase\s*(I{1,3}[ab\/]?\d*)', r'Phase\s*([1-4][ab]?)'])
        result['protocol_title']          = self._title(paragraphs, full_text)
        result['protocol_number']         = self._rx(full_text, [r'Protocol\s+(?:No\.?|Number|ID|#)\s*[:\-]?\s*([A-Za-z0-9][\w\-\/]{2,30})'])
        result['nct_number']              = self._rx(full_text, [r'(NCT\d{6,12})'])
        result['principal_investigator']  = self._near(full_text, ['Principal Investigator', 'Lead Investigator'])
        result['sponsor']                 = self._near(full_text, ['Sponsor', 'Sponsor Name'])
        result['funded_by']               = self._near(full_text, ['Funded by', 'Funding Source'])
        result['version_number']          = self._rx(full_text, [r'[Vv]ersion\s*([\d\.]+)', r'[Vv]\.\s*([\d\.]+)']) or 'v1.0'
        result['protocol_date']           = self._rx(full_text, [r'(?:Protocol\s+)?Date[:\s]+(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})', r'(\d{4}-\d{2}-\d{2})'])

        result['approval_data']['details'].update({
            'protocol_name': result['protocol_title'],
            'imp': self._near(full_text, ['IMP', 'Investigational Medicinal Product', 'Study Drug']),
            'indication': self._near(full_text, ['Indication', 'Condition', 'Disease']),
            'clinical_phase': phase,
            'investigators': result['principal_investigator'],
            'coordinating_investigator': self._near(full_text, ['Coordinating Investigator']),
            'expert_committee': self._near(full_text, ['Ethics Committee', 'IRB', 'IEC']),
            'sponsor_name_address': result['sponsor'],
        })

        def bl(patterns):
            for p in patterns:
                m = re.search(p, full_text, re.IGNORECASE | re.DOTALL)
                if m:
                    raw = full_text[m.end(): m.end() + 2000]
                    stop = re.search(r'\n(?:[A-Z\d][A-Z\s]{4,}|\d+\.[\s])', raw)
                    if stop: raw = raw[:stop.start()]
                    lines = [re.sub(r'^[\s\u2022\u00b7\-\*\d\.]+', '', l).strip() for l in raw.split('\n')]
                    res = [l for l in lines if 5 < len(l) < 400][:25]
                    if res: return res
            return []

        incl = bl([r'Inclusion\s+Criteria\s*[:\n]'])
        excl = bl([r'Exclusion\s+Criteria\s*[:\n]'])
        result['synopsis_data'].update({
            'overview': {
                'title': result['protocol_title'],
                'coordinating_investigator': self._near(full_text, ['Coordinating Investigator']),
                'expert_committee': self._near(full_text, ['Ethics Committee', 'IRB']),
                'investigators': result['principal_investigator'],
                'trial_sites': self._near(full_text, ['Study Site', 'Clinical Site']),
                'planned_period': self._near(full_text, ['Study Duration', 'Planned Period']),
                'fpfv': self._rx(full_text, [r'FPFV[:\s]+([^\n]+)']),
                'lplv': self._rx(full_text, [r'LPLV[:\s]+([^\n]+)']),
                'clinical_phase': phase,
            },
            'objectives': {
                'primary': bl([r'Primary\s+Objective[s]?\s*[:\n]']),
                'secondary': bl([r'Secondary\s+Objective[s]?\s*[:\n]']),
                'exploratory': bl([r'Exploratory\s+Objective[s]?\s*[:\n]']),
            },
            'endpoints': {
                'primary': bl([r'Primary\s+(?:Efficacy\s+)?Endpoint[s]?\s*[:\n]']),
                'secondary': bl([r'Secondary\s+(?:Efficacy\s+)?Endpoint[s]?\s*[:\n]']),
                'exploratory': bl([r'Exploratory\s+Endpoint[s]?\s*[:\n]']),
            },
            'num_patients': self._rx(full_text, [r'(?:Sample\s+[Ss]ize|[Nn]umber\s+of\s+(?:[Ss]ubjects|[Pp]atients))[:\s]*(\d+)', r'N\s*=\s*(\d+)']),
            'inclusion': {'text': '\n'.join(incl), 'points': incl},
            'exclusion': {'text': '\n'.join(excl), 'points': excl},
            'statistical_methods': '',
        })

        result['sections'] = sections
        result['soa_data'] = self._regex_soa(tables_data, soa_image_url)
        return result

    def _regex_soa(self, tables_data, soa_image_url) -> dict:
        CHECKS = {'x','✓','✔','•','y','yes','1','true','*','×'}
        SOA_KW = {'visit','screening','baseline','procedure','assessment','week','day','enrollment'}
        best, best_score = None, 0
        for tbl in tables_data:
            if len(tbl) < 2: continue
            score = sum(1 for kw in SOA_KW if kw in ' '.join(str(c) for c in tbl[0]).lower())
            if score > best_score:
                best_score, best = score, tbl
        if best and best_score >= 1:
            rh = [str(h).strip() for h in best[0]]
            if rh and rh[0].lower() not in ['procedure', 'assessment', 'activity', 'parameter']:
                headers, offset = ['Procedure'] + rh, True
            else:
                headers, offset = rh, False
            rows = []
            for row in best[1:]:
                cells = [str(c).strip() for c in row]
                if not any(cells): continue
                proc = '' if offset else (cells[0] if cells else '')
                cks  = cells if offset else cells[1:]
                processed = [proc] + ['1' if c.lower() in CHECKS else '0' for c in cks]
                while len(processed) < len(headers): processed.append('0')
                rows.append(processed[:len(headers)])
            if rows:
                return {'table': {'headers': headers, 'rows': rows}, 'image': None}
        if soa_image_url:
            return {'table': {'headers': [], 'rows': []},
                    'image': {'url': soa_image_url, 'caption': 'Schedule of Activities (from document)', 'description': ''}}
        return {'table': {'headers': [], 'rows': []}, 'image': None}

    def _all_tables_as_data(self, full_text):
        """Return stored tables from last parse run (populated by _assemble)."""
        return getattr(self, '_last_tables_data', [])

    # ──────────────────────────────────────────────
    # SMALL HELPERS
    # ──────────────────────────────────────────────
    def _title(self, paragraphs, full_text) -> str:
        v = self._rx(full_text, [
            r'(?:Full\s+)?(?:Protocol\s+)?Title[:\s]+([^\n]{10,200})',
            r'TITLE[:\s]+([^\n]{10,200})',
        ])
        if v: return v
        for p in paragraphs[:15]:
            if ('Heading' in p.get('style', '') or p.get('bold')) and len(p['text']) > 20:
                return p['text']
        return ''

    def _rx(self, text: str, patterns: List[str]) -> str:
        for pat in patterns:
            m = re.search(pat, text, re.IGNORECASE | re.MULTILINE)
            if m: return m.group(1).strip()
        return ''

    def _near(self, text: str, headings: List[str]) -> str:
        for h in headings:
            m = re.search(rf'(?:{re.escape(h)})\s*[:\.\-\n]\s*([^\n]{{3,300}})', text, re.IGNORECASE)
            if m: return m.group(1).strip()
        return ''

    def _save_image(self, img_bytes: bytes, ext: str, upload_dir: str) -> Optional[str]:
        try:
            os.makedirs(upload_dir, exist_ok=True)
            fname = f"{uuid.uuid4()}.{ext.lstrip('.').lower() or 'png'}"
            with open(os.path.join(upload_dir, fname), 'wb') as f:
                f.write(img_bytes)
            return f"/uploads/{fname}"
        except Exception as e:
            logger.warning(f"Image save: {e}")
            return None

    def _empty_protocol(self) -> dict:
        return {
            'protocol_title': '', 'protocol_number': '', 'nct_number': '',
            'principal_investigator': '', 'sponsor': '', 'funded_by': '',
            'version_number': 'v1.0', 'protocol_date': '',
            'sections': {}, 'synopsis': {}, 'schema_data': {'images': []},
            'approval_data': {
                'details': {
                    'protocol_name': '', 'protocol_number': '', 'imp': '', 'indication': '',
                    'clinical_phase': '', 'investigators': '', 'coordinating_investigator': '',
                    'expert_committee': '', 'sponsor_name_address': '',
                    'gcp_statement': '', 'approval_statement': ''
                },
                'sponsor_reps': [], 'cro_reps': [],
                'investigator_agreement': {
                    'description': '', 'signature': None, 'name': '', 'title': '',
                    'facility': '', 'city': '', 'state': '', 'date': ''
                },
                'amendments': []
            },
            'synopsis_data': {
                'overview': {
                    'title': '', 'coordinating_investigator': '', 'expert_committee': '',
                    'investigators': '', 'trial_sites': '', 'planned_period': '',
                    'fpfv': '', 'lplv': '', 'clinical_phase': ''
                },
                'objectives':  {'primary': [], 'secondary': [], 'exploratory': []},
                'endpoints':   {'primary': [], 'secondary': [], 'exploratory': []},
                'flowcharts': [], 'num_patients': '',
                'inclusion':   {'text': '', 'points': []},
                'exclusion':   {'text': '', 'points': []},
                'team': {'investigator_desc': '', 'coordinator_desc': ''},
                'tables': [], 'statistical_methods': ''
            },
            'section3': {
                'description': '',
                'image': {'url': None, 'caption': '', 'description': ''},
                'table': {'headers': ['Type', 'Objectives', 'Endpoints'], 'rows': []}
            },
            'soa_data': {'table': {'headers': [], 'rows': []}, 'image': None},
            'objectives_endpoints': [], 'abbreviations': [],
            'amendment_history': [], 'appendices': []
        }


document_parser = DocumentParser()
