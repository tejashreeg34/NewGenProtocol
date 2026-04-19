"""
word_generator.py
Generates a complete Clinical Trial Protocol Word (.docx) document
from the data submitted by the frontend, matching the Prot_1 reference structure.
"""

import os
import re
import io
import base64
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from PIL import Image as PILImage


# ============================================================================
# HELPERS
# ============================================================================

def strip_html(text):
    """Remove HTML tags from a string."""
    if not isinstance(text, str):
        return str(text) if text is not None else ''
    return re.sub(r'<[^>]+>', '', text).strip()


def safe(val, fallback=''):
    """Return stripped string or fallback if val is falsy."""
    v = strip_html(str(val)) if val is not None else ''
    return v if v else fallback


def get_image_bytes_from_data_url(data_url):
    """Convert a base64 data URL to a PIL-processed BytesIO stream."""
    if not data_url or not isinstance(data_url, str):
        return None
    if data_url.startswith('data:image/'):
        try:
            _, encoded = data_url.split(',', 1)
            raw = base64.b64decode(encoded)
            stream = io.BytesIO(raw)
            pil = PILImage.open(stream)
            if pil.mode == 'RGBA':
                bg = PILImage.new('RGB', pil.size, (255, 255, 255))
                bg.paste(pil, mask=pil.split()[3])
                pil = bg
            out = io.BytesIO()
            pil.save(out, format='PNG')
            out.seek(0)
            return out
        except Exception:
            return None
    return None


# ============================================================================
# STYLE SETUP
# ============================================================================

def set_custom_styles(doc):
    """Set up custom styles matching Prot_1 formatting."""
    styles = doc.styles

    def get_or_create(name, kind=WD_STYLE_TYPE.PARAGRAPH):
        return styles[name] if name in styles else styles.add_style(name, kind)

    # Protocol Title
    s = get_or_create('ProtocolTitle')
    s.font.name = 'Calibri'
    s.font.size = Pt(16)
    s.font.bold = True
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_after = Pt(24)

    # Heading 1
    h1 = get_or_create('Heading 1')
    h1.font.name = 'Calibri'
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True

    # Heading 2
    h2 = get_or_create('Heading 2')
    h2.font.name = 'Calibri'
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)

    # Heading 3
    h3 = get_or_create('Heading 3')
    h3.font.name = 'Calibri'
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(2)

    # TOC Heading (not included in TOC itself)
    toc_h = get_or_create('TOCHeading')
    toc_h.font.name = 'Calibri'
    toc_h.font.size = Pt(14)
    toc_h.font.bold = True
    toc_h.paragraph_format.space_before = Pt(12)
    toc_h.paragraph_format.space_after = Pt(6)

    # Normal Text
    nt = get_or_create('NormalText')
    nt.font.name = 'Calibri'
    nt.font.size = Pt(11)
    nt.paragraph_format.line_spacing = 1.15
    nt.paragraph_format.space_after = Pt(6)

    # Bullet List
    bl = get_or_create('BulletList')
    bl.font.name = 'Calibri'
    bl.font.size = Pt(11)
    bl.paragraph_format.left_indent = Pt(72)
    bl.paragraph_format.first_line_indent = Pt(-18)
    bl.paragraph_format.space_after = Pt(3)

    # Table Text
    tt = get_or_create('TableText')
    tt.font.name = 'Calibri'
    tt.font.size = Pt(9)
    tt.paragraph_format.space_after = Pt(2)


# ============================================================================
# HEADER / FOOTER
# ============================================================================

def add_header_footer(doc, pd):
    raw = pd.get('protocol_date', datetime.now().strftime('%Y-%m-%d'))
    try:
        formatted_date = datetime.strptime(raw, '%Y-%m-%d').strftime('%d %b %Y')
    except ValueError:
        formatted_date = raw

    version = pd.get('version_number', '1.0')
    title = pd.get('protocol_title', 'Clinical Trial Protocol')

    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(2.54)
        sec.right_margin = Cm(2.54)
        sec.header_distance = Cm(1.27)
        sec.footer_distance = Cm(1.27)

        # --- HEADER ---
        header = sec.header
        for p in header.paragraphs:
            p.text = ''

        tbl = header.add_table(rows=1, cols=2, width=Cm(16))
        tbl.autofit = True
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        tbl_borders = OxmlElement('w:tblBorders')
        for bn in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = OxmlElement(f'w:{bn}')
            b.set(qn('w:val'), 'nil')
            tbl_borders.append(b)
        tbl._tbl.tblPr.append(tbl_borders)

        cl = tbl.cell(0, 0).paragraphs[0]
        cl.text = title
        cl.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cl.style = doc.styles['Normal']

        cr = tbl.cell(0, 1).paragraphs[0]
        cr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cr.style = doc.styles['Normal']
        cr.add_run(f'Version {version}').add_break()
        cr.add_run(formatted_date)

        # --- FOOTER ---
        footer = sec.footer
        for p in footer.paragraphs:
            p.text = ''

        p_f1 = footer.add_paragraph(f'{title} – Version {version} {formatted_date}')
        p_f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_f1.style = doc.styles['Normal']
        p_f1.paragraph_format.space_after = Pt(0)

        p_f2 = footer.add_paragraph()
        p_f2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_f2.style = doc.styles['Normal']

        run_page = p_f2.add_run()
        fc1 = OxmlElement('w:fldChar')
        fc1.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = 'PAGE'
        fc2 = OxmlElement('w:fldChar')
        fc2.set(qn('w:fldCharType'), 'end')
        run_page._r.append(fc1)
        run_page._r.append(instr)
        run_page._r.append(fc2)


# ============================================================================
# TOC FIELD
# ============================================================================

def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fc = OxmlElement('w:fldChar')
    fc.set(qn('w:fldCharType'), 'begin')
    fc.set(qn('w:dirty'), 'true')
    run._r.append(fc)

    run2 = p.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    run2._r.append(instr)

    run3 = p.add_run()
    fc3 = OxmlElement('w:fldChar')
    fc3.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fc3)

    p.add_run('Right-click here and select "Update Field" to refresh the Table of Contents.')

    run4 = p.add_run()
    fc4 = OxmlElement('w:fldChar')
    fc4.set(qn('w:fldCharType'), 'end')
    run4._r.append(fc4)


# ============================================================================
# PROTOCOL INFO TABLE (Title Page)
# ============================================================================

def add_title_page(doc, pd):
    title = safe(pd.get('protocol_title'), 'Clinical Trial Protocol')
    doc.add_paragraph(title, style='ProtocolTitle')

    rows = [
        ('Protocol Number:', pd.get('protocol_number', '')),
        ('National Clinical Trial (NCT) Identified Number:', pd.get('nct_number', '')),
        ('Principal Investigator:', pd.get('principal_investigator', '')),
        ('IND/IDE Sponsor:', pd.get('sponsor', '')),
        ('Funded by:', pd.get('funded_by', '')),
        ('Version Number:', pd.get('version_number', '')),
        ('Date:', pd.get('protocol_date', datetime.now().strftime('%d %B %Y'))),
    ]

    doc.add_paragraph()
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    table.autofit = False

    for row in table.rows:
        row.cells[0].width = Cm(7)
        row.cells[1].width = Cm(9)

    for i, (label, value) in enumerate(rows):
        c0 = table.cell(i, 0)
        c0.text = label
        c0.paragraphs[0].style = 'TableText'
        c0.paragraphs[0].runs[0].bold = True

        c1 = table.cell(i, 1)
        c1.text = safe(value)       # blank if empty — row still shows
        c1.paragraphs[0].style = 'TableText'

    # Summary of Changes (optional)
    summary = safe(pd.get('summary_changes', ''))
    if summary:
        doc.add_paragraph()
        p = doc.add_paragraph(style='NormalText')
        p.add_run('Summary of Changes from Previous Version:').bold = True
        doc.add_paragraph(summary, style='NormalText')


# ============================================================================
# PROTOCOL APPROVAL & AGREEMENT
# ============================================================================

def add_approval_section(doc, pd):
    approval = pd.get('approval_data') or {}
    details = approval.get('details') or {}

    # Only render this section if there is any data at all
    has_detail = any(safe(v) for v in details.values())
    has_reps = bool(approval.get('sponsor_reps') or approval.get('cro_reps'))
    has_agree = bool(approval.get('investigator_agreement'))

    if not (has_detail or has_reps or has_agree):
        return

    doc.add_heading('PROTOCOL APPROVAL & AGREEMENT', level=1)

    # Details table
    detail_rows = [
        ('Protocol Name:', details.get('protocol_name', '')),
        ('Protocol Number:', details.get('protocol_number', '')),
        ('IMP:', details.get('imp', '')),
        ('Indication:', details.get('indication', '')),
        ('Clinical Phase:', details.get('clinical_phase', '')),
        ('Investigators:', details.get('investigators', '')),
        ('Coordinating Investigator:', details.get('coordinating_investigator', '')),
        ('Expert Committee:', details.get('expert_committee', '')),
        ('Sponsor Name & Address:', details.get('sponsor_name_address', '')),
    ]

    table = doc.add_table(rows=len(detail_rows), cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Cm(7)
        row.cells[1].width = Cm(9)

    for i, (label, value) in enumerate(detail_rows):
        c0 = table.cell(i, 0)
        c0.text = label
        c0.paragraphs[0].style = 'TableText'
        c0.paragraphs[0].runs[0].bold = True
        c1 = table.cell(i, 1)
        c1.text = safe(value)
        c1.paragraphs[0].style = 'TableText'

    doc.add_paragraph()

    gcp = safe(details.get('gcp_statement', ''))
    if gcp:
        p = doc.add_paragraph(style='NormalText')
        p.add_run('GCP Statement').bold = True
        doc.add_paragraph(gcp, style='NormalText')

    approval_stmt = safe(details.get('approval_statement', ''))
    if approval_stmt:
        p = doc.add_paragraph(style='NormalText')
        p.add_run('Approval Statement').bold = True
        doc.add_paragraph(approval_stmt, style='NormalText')

    def render_rep(rep, role_label):
        doc.add_paragraph(role_label, style='Heading 3')
        rep_rows = [
            ('Name:', rep.get('name', '')),
            ('Title:', rep.get('title', '')),
            ('Organization:', rep.get('organization', '')),
            ('Date:', rep.get('date', '')),
        ]
        rt = doc.add_table(rows=len(rep_rows), cols=2)
        rt.style = 'Table Grid'
        rt.autofit = False
        for row in rt.rows:
            row.cells[0].width = Cm(5)
            row.cells[1].width = Cm(11)
        for j, (lbl, val) in enumerate(rep_rows):
            c0 = rt.cell(j, 0)
            c0.text = lbl
            c0.paragraphs[0].style = 'TableText'
            c0.paragraphs[0].runs[0].bold = True
            c1 = rt.cell(j, 1)
            c1.text = safe(val)
            c1.paragraphs[0].style = 'TableText'

        doc.add_paragraph()
        sig_url = rep.get('signature')
        sig_stream = get_image_bytes_from_data_url(sig_url)
        p_sig = doc.add_paragraph(style='NormalText')
        p_sig.add_run('Signature: ').bold = True
        if sig_stream:
            try:
                doc.add_picture(sig_stream, width=Cm(5))
            except Exception:
                p_sig.add_run('____________________')
        else:
            p_sig.add_run('____________________')
        doc.add_paragraph()

    for rep in (approval.get('sponsor_reps') or []):
        render_rep(rep, 'Sponsor Representative')

    for rep in (approval.get('cro_reps') or []):
        render_rep(rep, 'CRO Representative')

    agree = approval.get('investigator_agreement') or {}
    if any(safe(v) for k, v in agree.items() if k != 'signature'):
        doc.add_heading('Investigator Agreement', level=2)

        desc = safe(agree.get('description', ''))
        if desc:
            doc.add_paragraph(desc, style='NormalText')

        agree_rows = [
            ('Investigator Name:', agree.get('name', '')),
            ('Title:', agree.get('title', '')),
            ('Facility:', agree.get('facility', '')),
            ('City:', agree.get('city', '')),
            ('State:', agree.get('state', '')),
            ('Date:', agree.get('date', '')),
        ]
        at = doc.add_table(rows=len(agree_rows), cols=2)
        at.style = 'Table Grid'
        at.autofit = False
        for row in at.rows:
            row.cells[0].width = Cm(5)
            row.cells[1].width = Cm(11)
        for j, (lbl, val) in enumerate(agree_rows):
            c0 = at.cell(j, 0)
            c0.text = lbl
            c0.paragraphs[0].style = 'TableText'
            c0.paragraphs[0].runs[0].bold = True
            c1 = at.cell(j, 1)
            c1.text = safe(val)
            c1.paragraphs[0].style = 'TableText'

        doc.add_paragraph()
        sig_stream = get_image_bytes_from_data_url(agree.get('signature'))
        p_sig = doc.add_paragraph(style='NormalText')
        p_sig.add_run('Signature: ').bold = True
        if sig_stream:
            try:
                doc.add_picture(sig_stream, width=Cm(5))
            except Exception:
                p_sig.add_run('____________________')
        else:
            p_sig.add_run('____________________')

    doc.add_page_break()


# ============================================================================
# SECTION 1: PROTOCOL SUMMARY
# ============================================================================

def _render_list_items(doc, items):
    """Render a list of items as bullet points."""
    for item in items:
        text = safe(item) if not isinstance(item, dict) else safe(item.get('text', str(item)))
        if text:
            p = doc.add_paragraph(style='BulletList')
            p.add_run(f'• {text}')


def add_synopsis_section(doc, pd):
    """1.1 Synopsis — uses synopsis_data.overview and synopsis_data.objectives/endpoints."""
    s_data = pd.get('synopsis_data') or {}
    overview = s_data.get('overview') or {}
    objectives = s_data.get('objectives') or {}
    endpoints = s_data.get('endpoints') or {}
    inclusion = s_data.get('inclusion') or {}
    exclusion = s_data.get('exclusion') or {}
    team = s_data.get('team') or {}

    doc.add_heading('1.1 Synopsis', level=2)

    # Overview table
    ov_rows = [
        ('Title:', overview.get('title', '')),
        ('Clinical Phase:', overview.get('clinical_phase', '')),
        ('Coordinating Investigator:', overview.get('coordinating_investigator', '')),
        ('Expert Committee:', overview.get('expert_committee', '')),
        ('Investigators:', overview.get('investigators', '')),
        ('Trial Sites:', overview.get('trial_sites', '')),
        ('Planned Study Period:', overview.get('planned_period', '')),
        ('First Patient First Visit (FPFV):', overview.get('fpfv', '')),
        ('Last Patient Last Visit (LPLV):', overview.get('lplv', '')),
        ('Number of Patients:', s_data.get('num_patients', '')),
    ]

    tbl = doc.add_table(rows=len(ov_rows), cols=2)
    tbl.style = 'Table Grid'
    tbl.autofit = False
    for row in tbl.rows:
        row.cells[0].width = Cm(7)
        row.cells[1].width = Cm(9)
    for i, (label, value) in enumerate(ov_rows):
        c0 = tbl.cell(i, 0)
        c0.text = label
        c0.paragraphs[0].style = 'TableText'
        c0.paragraphs[0].runs[0].bold = True
        c1 = tbl.cell(i, 1)
        c1.text = safe(value)
        c1.paragraphs[0].style = 'TableText'

    doc.add_paragraph()

    # Objectives
    def _safe_list(lst):
        return [x for x in (lst or []) if x]

    if _safe_list(objectives.get('primary', [])):
        doc.add_paragraph('Primary Objectives', style='Heading 3')
        _render_list_items(doc, objectives['primary'])
    if _safe_list(objectives.get('secondary', [])):
        doc.add_paragraph('Secondary Objectives', style='Heading 3')
        _render_list_items(doc, objectives['secondary'])
    if _safe_list(objectives.get('exploratory', [])):
        doc.add_paragraph('Exploratory Objectives', style='Heading 3')
        _render_list_items(doc, objectives['exploratory'])

    # Endpoints
    if _safe_list(endpoints.get('primary', [])):
        doc.add_paragraph('Primary Endpoints', style='Heading 3')
        _render_list_items(doc, endpoints['primary'])
    if _safe_list(endpoints.get('secondary', [])):
        doc.add_paragraph('Secondary Endpoints', style='Heading 3')
        _render_list_items(doc, endpoints['secondary'])
    if _safe_list(endpoints.get('exploratory', [])):
        doc.add_paragraph('Exploratory Endpoints', style='Heading 3')
        _render_list_items(doc, endpoints['exploratory'])

    # Inclusion Criteria
    inc_text = safe(inclusion.get('text', ''))
    inc_points = _safe_list(inclusion.get('points', []))
    if inc_text or inc_points:
        doc.add_paragraph('Inclusion Criteria', style='Heading 3')
        if inc_text:
            doc.add_paragraph(inc_text, style='NormalText')
        _render_list_items(doc, inc_points)

    # Exclusion Criteria
    exc_text = safe(exclusion.get('text', ''))
    exc_points = _safe_list(exclusion.get('points', []))
    if exc_text or exc_points:
        doc.add_paragraph('Exclusion Criteria', style='Heading 3')
        if exc_text:
            doc.add_paragraph(exc_text, style='NormalText')
        _render_list_items(doc, exc_points)

    # Statistical Methods
    stat = safe(s_data.get('statistical_methods', ''))
    if stat:
        doc.add_paragraph('Statistical Methods', style='Heading 3')
        doc.add_paragraph(stat, style='NormalText')

    # Team
    inv_desc = safe(team.get('investigator_desc', ''))
    if inv_desc:
        doc.add_paragraph('Investigators', style='Heading 3')
        doc.add_paragraph(inv_desc, style='NormalText')
    coord_desc = safe(team.get('coordinator_desc', ''))
    if coord_desc:
        doc.add_paragraph('Coordinating Investigator', style='Heading 3')
        doc.add_paragraph(coord_desc, style='NormalText')

    # Flowcharts
    flowcharts = s_data.get('flowcharts') or []
    if flowcharts:
        fc_title = safe(s_data.get('flowchart_title', 'Study Schema / Flowchart'))
        doc.add_paragraph(fc_title, style='Heading 3')
        fc_desc = safe(s_data.get('flowchart_description', ''))
        if fc_desc:
            doc.add_paragraph(fc_desc, style='NormalText')
        for fc in flowcharts:
            url = fc.get('url') if isinstance(fc, dict) else fc
            stream = get_image_bytes_from_data_url(url)
            if stream:
                try:
                    doc.add_picture(stream, width=Cm(15))
                except Exception:
                    pass
            elif url and isinstance(url, str) and os.path.exists(url.lstrip('/')):
                try:
                    doc.add_picture(url.lstrip('/'), width=Cm(15))
                except Exception:
                    pass
            cap = fc.get('caption', '') if isinstance(fc, dict) else ''
            if cap:
                p = doc.add_paragraph(style='NormalText')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(strip_html(cap)).bold = True

    # Custom summary tables
    tables = s_data.get('tables') or []
    for t in tables:
        headers = t.get('headers', [])
        rows = t.get('rows', [])
        if not headers:
            continue
        title_t = safe(t.get('title', ''))
        if title_t:
            doc.add_paragraph(title_t, style='Heading 3')
        add_dynamic_table(doc, headers, rows)


def add_schema_section(doc, pd):
    """1.2 Schema"""
    schema = pd.get('schema_data') or {}
    images = schema.get('images') or []
    if not images and schema.get('image_url'):
        images = [{'url': schema['image_url'], 'caption': schema.get('caption', ''), 'description': schema.get('description', '')}]

    if not images:
        return

    doc.add_heading('1.2 Schema', level=2)
    for img in images:
        url = img.get('url') if isinstance(img, dict) else img
        stream = get_image_bytes_from_data_url(url)
        if stream:
            try:
                doc.add_picture(stream, width=Cm(15))
            except Exception:
                pass
        elif url and isinstance(url, str):
            path = url.lstrip('/')
            if os.path.exists(path):
                try:
                    doc.add_picture(path, width=Cm(15))
                except Exception:
                    pass
        cap = img.get('caption', '') if isinstance(img, dict) else ''
        desc = img.get('description', '') if isinstance(img, dict) else ''
        if safe(cap):
            p = doc.add_paragraph(style='NormalText')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(safe(cap)).bold = True
        if safe(desc):
            doc.add_paragraph(safe(desc), style='NormalText')
    doc.add_paragraph()


def add_soa_section(doc, pd):
    """1.3 Schedule of Activities (SoA)"""
    soa = pd.get('soa_data') or {}

    has_image = bool((soa.get('image') or {}).get('url'))
    has_table = bool((soa.get('table') or {}).get('headers'))

    if not has_image and not has_table:
        return

    doc.add_heading('1.3 Schedule of Activities (SoA)', level=2)

    # SoA image (if uploaded)
    soa_img = soa.get('image') or {}
    if soa_img.get('url'):
        stream = get_image_bytes_from_data_url(soa_img['url'])
        if stream:
            try:
                doc.add_picture(stream, width=Cm(15))
            except Exception:
                pass
        elif os.path.exists(soa_img['url'].lstrip('/')):
            try:
                doc.add_picture(soa_img['url'].lstrip('/'), width=Cm(15))
            except Exception:
                pass
        if safe(soa_img.get('caption', '')):
            p = doc.add_paragraph(style='NormalText')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(safe(soa_img['caption'])).bold = True
        if safe(soa_img.get('description', '')):
            doc.add_paragraph(safe(soa_img['description']), style='NormalText')

    # SoA grid table
    tbl_data = soa.get('table') or {}
    headers = tbl_data.get('headers', [])
    rows = tbl_data.get('rows', {})
    if headers:
        col_headers = ['Procedures'] + list(headers)
        table = doc.add_table(rows=1, cols=len(col_headers))
        table.style = 'Table Grid'

        hdr = table.rows[0].cells
        for i, h in enumerate(col_headers):
            hdr[i].text = safe(str(h))
            hdr[i].paragraphs[0].style = 'TableText'
            hdr[i].paragraphs[0].runs[0].bold = True

        if isinstance(rows, dict):
            for proc, checks in rows.items():
                cells = table.add_row().cells
                cells[0].text = safe(str(proc))
                cells[0].paragraphs[0].style = 'TableText'
                for i, val in enumerate(checks):
                    cells[i + 1].text = 'X' if val else ''
                    cells[i + 1].paragraphs[0].style = 'TableText'
        elif isinstance(rows, list):
            for row in rows:
                cells = table.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = safe(str(val))
                    cells[i].paragraphs[0].style = 'TableText'

        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_paragraph()


# ============================================================================
# SECTION 3: OBJECTIVES & ENDPOINTS
# ============================================================================

def add_section3(doc, pd):
    """Section 3: Objectives and Endpoints"""
    s3 = pd.get('section3') or {}
    desc = safe(s3.get('description', ''))
    tbl = s3.get('table') or {}
    img = s3.get('image') or {}
    legacy_objs = pd.get('objectives_endpoints', [])

    # Defensive checks for dict types
    tbl_headers = []
    if isinstance(tbl, dict):
        tbl_headers = tbl.get('headers') or []
    
    img_url = None
    if isinstance(img, dict):
        img_url = img.get('url')
    elif isinstance(img, str):
        img_url = img

    if not desc and not tbl_headers and not img_url and not legacy_objs:
        return

    if desc:
        doc.add_paragraph(desc, style='NormalText')

    if img_url:
        stream = get_image_bytes_from_data_url(img_url)
        if stream:
            try:
                doc.add_picture(stream, width=Cm(15))
            except Exception:
                pass
        elif os.path.exists(img_url.lstrip('/')):
            try:
                doc.add_picture(img_url.lstrip('/'), width=Cm(15))
            except Exception:
                pass
        
        img_caption = img.get('caption', '') if isinstance(img, dict) else ''
        img_desc = img.get('description', '') if isinstance(img, dict) else ''
        
        if safe(img_caption):
            p = doc.add_paragraph(style='NormalText')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"Figure: {safe(img_caption)}").bold = True
        if safe(img_desc):
            doc.add_paragraph(safe(img_desc), style='NormalText')

    if tbl_headers:
        add_dynamic_table(doc, tbl_headers, tbl.get('rows', []) if isinstance(tbl, dict) else [])
    elif legacy_objs:
        add_objectives_table(doc, legacy_objs)


# ============================================================================
# TABLE HELPERS
# ============================================================================

def add_dynamic_table(doc, headers, rows):
    if not headers:
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = safe(str(h))
        hdr_cells[i].paragraphs[0].style = 'TableText'
        if hdr_cells[i].paragraphs[0].runs:
            hdr_cells[i].paragraphs[0].runs[0].bold = True

    for row_data in (rows or []):
        row_cells = table.add_row().cells
        if isinstance(row_data, list):
            for i in range(len(headers)):
                val = row_data[i] if i < len(row_data) else ''
                row_cells[i].text = safe(str(val))
                row_cells[i].paragraphs[0].style = 'TableText'
        elif isinstance(row_data, dict):
            for i, h in enumerate(headers):
                row_cells[i].text = safe(str(row_data.get(h, '')))
                row_cells[i].paragraphs[0].style = 'TableText'

    doc.add_paragraph()


def add_objectives_table(doc, objectives_data):
    if not objectives_data:
        return
    headers = ['Type', 'Objectives', 'Endpoints and Justification']
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].style = 'TableText'
        c.paragraphs[0].runs[0].bold = True

    for row_data in objectives_data:
        cells = table.add_row().cells
        cells[0].text = safe(row_data.get('Type', ''))
        cells[1].text = safe(row_data.get('Objective', ''))
        ep = safe(row_data.get('Endpoint', ''))
        just = safe(row_data.get('Justification', ''))
        cells[2].text = f"{ep}\n\nJustification: {just}" if just else ep
        for cell in cells:
            cell.paragraphs[0].style = 'TableText'
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    doc.add_paragraph()


def add_abbreviations_table(doc, abbreviations):
    if not abbreviations:
        return
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Abbreviation'
    hdr[1].text = 'Full Form'
    for cell in hdr:
        cell.paragraphs[0].style = 'TableText'
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True

    for ab in abbreviations:
        cells = table.add_row().cells
        cells[0].text = safe(ab.get('Abbreviation', ab.get('abbreviation', '')))
        cells[1].text = safe(ab.get('Full Form', ab.get('full_form', '')))
        for cell in cells:
            cell.paragraphs[0].style = 'TableText'

    doc.add_paragraph()


def add_amendment_table(doc, amendments):
    if not amendments:
        return
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    headers = ['Version', 'Date', 'Description of Change', 'Brief Rationale']
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].style = 'TableText'
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].bold = True

    for am in amendments:
        cells = table.add_row().cells
        cells[0].text = safe(am.get('Version', am.get('version', '')))
        cells[1].text = safe(am.get('Date', am.get('date', '')))
        cells[2].text = safe(am.get('Description', am.get('description', '')))
        cells[3].text = safe(am.get('Rationale', am.get('rationale', '')))
        for cell in cells:
            cell.paragraphs[0].style = 'TableText'
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    doc.add_paragraph()


# ============================================================================
# SECTION IMAGE HELPER
# ============================================================================

def _add_section_images(doc, images):
    for img in (images or []):
        url = img.get('url') if isinstance(img, dict) else img
        stream = get_image_bytes_from_data_url(url)
        if stream:
            try:
                doc.add_picture(stream, width=Cm(15))
            except Exception:
                pass
        elif url and isinstance(url, str) and os.path.exists(url.lstrip('/')):
            try:
                doc.add_picture(url.lstrip('/'), width=Cm(15))
            except Exception:
                pass
        cap = img.get('caption', '') if isinstance(img, dict) else ''
        desc = img.get('description', '') if isinstance(img, dict) else ''
        if safe(cap):
            p = doc.add_paragraph(style='NormalText')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(safe(cap)).bold = True
        if safe(desc):
            doc.add_paragraph(safe(desc), style='NormalText')


# ============================================================================
# TEMPLATE STRUCTURE
# ============================================================================

TEMPLATE_STRUCTURE = [
    {'id': 1,  'title': 'PROTOCOL SUMMARY',
     'subsections': ['Synopsis', 'Schema', 'Schedule of Activities (SoA)']},
    {'id': 2,  'title': 'INTRODUCTION',
     'subsections': ['Study Rationale', 'Background', 'Risk/Benefit Assessment',
                     'Known Potential Risks', 'Known Potential Benefits', 'Assessment of Potential Risks and Benefits']},
    {'id': 3,  'title': 'OBJECTIVES AND ENDPOINTS', 'subsections': []},
    {'id': 4,  'title': 'STUDY DESIGN',
     'subsections': ['Overall Design', 'Scientific Rationale for Study Design', 'Justification for Dose', 'End of Study Definition']},
    {'id': 5,  'title': 'STUDY POPULATION',
     'subsections': ['Inclusion Criteria', 'Exclusion Criteria', 'Lifestyle Considerations', 'Screen Failures', 'Strategies for Recruitment and Retention']},
    {'id': 6,  'title': 'STUDY INTERVENTION',
     'subsections': ['Study Intervention(s) Administration', 'Study Intervention Description', 'Dosing and Administration',
                     'Preparation/Handling/Storage/Accountability', 'Acquisition and accountability',
                     'Formulation, Appearance, Packaging, and Labeling', 'Product Storage and Stability',
                     'Preparation', 'Measures to Minimize Bias: Randomization and Blinding',
                     'Study Intervention Compliance', 'Concomitant Therapy', 'Rescue Medicine']},
    {'id': 7,  'title': 'STUDY INTERVENTION DISCONTINUATION AND PARTICIPANT DISCONTINUATION/WITHDRAWAL',
     'subsections': ['Discontinuation of Study Intervention', 'Participant Discontinuation/Withdrawal from the Study', 'Lost to Follow-Up']},
    {'id': 8,  'title': 'STUDY ASSESSMENTS AND PROCEDURES',
     'subsections': ['Efficacy Assessments', 'Safety and Other Assessments', 'Adverse Events and Serious Adverse Events',
                     'Definition of Adverse Events (AE)', 'Definition of Serious Adverse Events (SAE)',
                     'Classification of an Adverse Event', 'Time Period and Frequency for Event Assessment and Follow-Up',
                     'Adverse Event Reporting', 'Serious Adverse Event Reporting', 'Reporting Events to Participants',
                     'Events of Special Interest', 'Reporting of Pregnancy', 'Unanticipated Problems',
                     'Definition of Unanticipated Problems (UP)', 'Unanticipated Problem Reporting',
                     'Reporting Unanticipated Problems to Participants']},
    {'id': 9,  'title': 'STATISTICAL CONSIDERATIONS',
     'subsections': ['Statistical Hypotheses', 'Sample Size Determination', 'Populations for Analyses',
                     'Statistical Analyses', 'General Approach', 'Analysis of the Primary Efficacy Endpoint(s)',
                     'Analysis of the Secondary Endpoint(s)', 'Safety Analyses', 'Baseline Descriptive Statistics',
                     'Planned Interim Analyses', 'Sub-Group Analyses', 'Tabulation of Individual participant Data', 'Exploratory Analyses']},
    {'id': 10, 'title': 'SUPPORTING DOCUMENTATION AND OPERATIONAL CONSIDERATIONS',
     'subsections': ['Regulatory, Ethical, and Study Oversight Considerations', 'Informed Consent Process',
                     'Study Discontinuation and Closure', 'Confidentiality and Privacy',
                     'Future Use of Stored Specimens and Data', 'Key Roles and Study Governance',
                     'Safety Oversight', 'Clinical Monitoring', 'Quality Assurance and Quality Control',
                     'Data Handling and Record Keeping', 'Protocol Deviations', 'Publication and Data Sharing Policy',
                     'Conflict of Interest Policy', 'Additional Considerations', 'Abbreviations', 'Protocol Amendment History']},
    {'id': 11, 'title': 'REFERENCES', 'subsections': []},
]


# ============================================================================
# MAIN ENTRY POINT
# ============================================================================

def generate_complete_word_document(protocol_data):
    """Generate complete Word document matching the Prot_1 reference structure."""
    pd = protocol_data  # alias for brevity
    doc = Document()
    set_custom_styles(doc)
    add_header_footer(doc, pd)

    # ═══ TITLE PAGE ═══
    add_title_page(doc, pd)
    doc.add_page_break()

    # ═══ PROTOCOL APPROVAL & AGREEMENT ═══
    add_approval_section(doc, pd)

    # ═══ STATEMENT OF COMPLIANCE ═══
    doc.add_heading('STATEMENT OF COMPLIANCE', level=1)
    compliance = (
        "The trial will be carried out in accordance with International Conference on Harmonisation Good "
        "Clinical Practice (ICH GCP) and the following:\n\n"
        "• United States (US) Code of Federal Regulations (CFR) applicable to clinical studies "
        "(45 CFR Part 46, 21 CFR Part 50, 21 CFR Part 56, 21 CFR Part 312, and/or 21 CFR Part 812)\n\n"
        "• National Institutes of Health (NIH)-funded investigators and clinical trial site staff who are "
        "responsible for the conduct, management, or oversight of NIH-funded clinical trials have completed "
        "Human Subjects Protection and ICH GCP Training.\n\n"
        "• The protocol, informed consent form(s), recruitment materials, and all participant materials will "
        "be submitted to the Institutional Review Board (IRB) for review and approval. Approval of both the "
        "protocol and the consent form must be obtained before any participant is enrolled. Any amendment to "
        "the protocol will require review and approval by the IRB before the changes are implemented to the study."
    )
    doc.add_paragraph(compliance, style='NormalText')
    doc.add_page_break()

    # ═══ TABLE OF CONTENTS ═══
    doc.add_paragraph('Table of Contents', style='TOCHeading')
    add_toc(doc)
    doc.add_page_break()

    # ═══ SECTION 1: PROTOCOL SUMMARY ═══
    doc.add_heading('1 PROTOCOL SUMMARY', level=1)
    add_synopsis_section(doc, pd)
    add_schema_section(doc, pd)
    add_soa_section(doc, pd)
    doc.add_page_break()

    # ═══ SECTION 2: INTRODUCTION ═══
    doc.add_heading('2 INTRODUCTION', level=1)
    sections = pd.get('sections') or {}
    sec2 = sections.get('2') or {}
    if safe(sec2.get('main', '')):
        doc.add_paragraph(safe(sec2['main']), style='NormalText')

    s2_template = next((s for s in TEMPLATE_STRUCTURE if s['id'] == 2), None)
    s2_subs = sec2.get('subsections') or []

    if isinstance(s2_subs, list):
        for i, sub in enumerate(s2_subs):
            if isinstance(sub, dict):
                title = safe(sub.get('title', ''))
                content = safe(sub.get('content', ''))
                heading_num = f"2.{i + 1} {title}"
                doc.add_heading(heading_num, level=2)
                if content:
                    doc.add_paragraph(content, style='NormalText')
                # Subsection images
                _add_section_images(doc, sub.get('images') or [])
    elif isinstance(s2_subs, dict):
        for idx, content in sorted(s2_subs.items(), key=lambda x: int(x[0])):
            if content:
                doc.add_heading(f"2.{int(idx)+1} Subsection", level=2)
                doc.add_paragraph(safe(str(content)), style='NormalText')

    doc.add_page_break()

    # ═══ SECTION 3: OBJECTIVES AND ENDPOINTS ═══
    doc.add_heading('3 OBJECTIVES AND ENDPOINTS', level=1)
    add_section3(doc, pd)
    doc.add_page_break()

    # ═══ SECTIONS 4–11 ═══
    for sec_num in range(4, 12):
        sec_key = str(sec_num)
        template = next((s for s in TEMPLATE_STRUCTURE if s['id'] == sec_num), None)
        sec_title = f"{sec_num} {template['title']}" if template else f"{sec_num} SECTION {sec_num}"
        doc.add_heading(sec_title, level=1)

        sec_data = sections.get(sec_key) or {}
        main_text = safe(sec_data.get('main', ''))
        if main_text:
            doc.add_paragraph(main_text, style='NormalText')

        # Section-level images
        _add_section_images(doc, sec_data.get('images') or [])

        subs = sec_data.get('subsections') or []

        if isinstance(subs, list):
            for i, sub in enumerate(subs):
                if isinstance(sub, dict):
                    sub_title = safe(sub.get('title', ''))
                    sub_content = safe(sub.get('content', ''))
                    heading_num = f"{sec_num}.{i + 1} {sub_title}"
                    doc.add_heading(heading_num, level=2)
                    if sub_content:
                        doc.add_paragraph(sub_content, style='NormalText')
                    _add_section_images(doc, sub.get('images') or [])

                    # Custom table inside subsection
                    custom_tbl = sub.get('customTable') or {}
                    if custom_tbl.get('headers'):
                        add_dynamic_table(doc, custom_tbl['headers'], custom_tbl.get('rows', []))

                elif isinstance(sub, str) and sub.strip():
                    doc.add_heading(f"{sec_num}.{i + 1} {sub}", level=2)

        elif isinstance(subs, dict):
            for idx, content in sorted(subs.items(), key=lambda x: int(x[0])):
                if content:
                    doc.add_heading(f"{sec_num}.{int(idx)+1} Subsection", level=2)
                    doc.add_paragraph(safe(str(content)), style='NormalText')

        # Special handling for Section 10: Abbreviations & Amendment History
        if sec_num == 10:
            abbrevs = pd.get('abbreviations') or sec_data.get('abbreviations') or []
            if abbrevs:
                doc.add_heading(f'10.{len(subs)+1} Abbreviations', level=2)
                add_abbreviations_table(doc, abbrevs)

            amendments = pd.get('amendment_history') or sec_data.get('amendment_history') or []
            if amendments:
                doc.add_heading(f'10.{len(subs)+2} Protocol Amendment History', level=2)
                add_amendment_table(doc, amendments)

        doc.add_page_break()

    # ═══ APPENDICES ═══
    appendices = pd.get('appendices') or []
    if appendices:
        doc.add_heading('APPENDICES', level=1)
        for i, appendix in enumerate(appendices, 1):
            doc.add_heading(f'Appendix {i}', level=2)
            app_title = safe(appendix.get('title', ''))
            app_content = safe(appendix.get('content', ''))
            if app_title:
                doc.add_paragraph(app_title, style='NormalText')
            if app_content:
                doc.add_paragraph(app_content, style='NormalText')

    # ═══ SAVE ═══
    output_dir = 'generated_docs'
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    version = safe(pd.get('version_number', 'v1.0')).replace('.', '_').replace(' ', '_')
    filename = f'protocol_{version}_{timestamp}.docx'
    filepath = os.path.join(output_dir, filename)

    # Force TOC update on open
    elem = OxmlElement('w:updateFields')
    elem.set(qn('w:val'), 'true')
    doc.settings.element.append(elem)

    doc.save(filepath)
    return filepath
