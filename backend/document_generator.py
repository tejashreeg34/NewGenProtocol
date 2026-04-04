from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from datetime import datetime
import re

def add_html_to_paragraph(paragraph, html_text):
    """
    Parses a simple HTML string and adds formatted runs to the paragraph.
    Supported tags: <b>, <strong>, <i>, <em>, <u>, <br>, <div>, <p>
    """
    if not html_text:
        return
    
    # Simple regex-based parser for basic tags
    # We split by tags and keep them in the resulting list
    parts = re.split(r'(<[^>]+>)', str(html_text))
    
    current_styles = {
        'bold': False,
        'italic': False,
        'underline': False
    }
    
    for part in parts:
        if not part:
            continue
        
        tag_match = re.match(r'<(/?)([^ >]+)', part.lower())
        if tag_match:
            is_closing = tag_match.group(1) == '/'
            tag_name = tag_match.group(2)
            
            if tag_name in ['b', 'strong']:
                current_styles['bold'] = not is_closing
            elif tag_name in ['i', 'em']:
                current_styles['italic'] = not is_closing
            elif tag_name in ['u']:
                current_styles['underline'] = not is_closing
            elif tag_name in ['br', 'div', 'p']:
                # For block elements, we just add a break if it's not the start of the tag
                if tag_name == 'br' or (not is_closing and tag_name in ['div', 'p']):
                    paragraph.add_run().add_break()
        else:
            # It's text
            # Decode common HTML entities
            text = part.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
            if text:
                run = paragraph.add_run(text)
                run.bold = current_styles['bold']
                run.italic = current_styles['italic']
                run.underline = current_styles['underline']

def add_labeled_paragraph(doc, label, value, style='NormalText', bold_label=True):
    """Add a paragraph with a bold label and normal value"""
    p = doc.add_paragraph(style=style)
    # Strip any tags from label for safety as it's usually just a string
    label = re.sub(r'<.*?>', '', str(label))
    
    run_label = p.add_run(label)
    if bold_label:
        run_label.bold = True
    
    p.add_run(" ")
    add_html_to_paragraph(p, value)
    return p

import os
import io
import json
import base64

# Define the standard structure for sections
TEMPLATE_STRUCTURE = [
    {
        "id": 1,
        "title": "PROTOCOL SUMMARY",
        "subsections": [
            "Synopsis",
            "Schema",
            "Schedule of Activities (SoA)"
        ]
    },
    {
        "id": 2,
        "title": "INTRODUCTION",
        "subsections": [
            "Study Rationale",
            "Background",
            "Risk/Benefit Assessment",
            "Known Potential Risks",
            "Known Potential Benefits",
            "Assessment of Potential Risks and Benefits"
        ]
    },
    {
        "id": 3,
        "title": "OBJECTIVES AND ENDPOINTS",
        "subsections": []
    },
    {
        "id": 4,
        "title": "STUDY DESIGN",
        "subsections": [
            "Overall Design",
            "Scientific Rationale for Study Design",
            "Justification for Dose",
            "End of Study Definition"
        ]
    },
    {
        "id": 5,
        "title": "STUDY POPULATION",
        "subsections": [
            "Inclusion Criteria",
            "Exclusion Criteria",
            "Lifestyle Considerations",
            "Screen Failures",
            "Strategies for Recruitment and Retention"
        ]
    },
    {
        "id": 6,
        "title": "STUDY INTERVENTION",
        "subsections": [
            "Study Intervention(s) Administration",
            "Study Intervention Description",
            "Dosing and Administration",
            "Preparation/Handling/Storage/Accountability",
            "Acquisition and accountability",
            "Formulation, Appearance, Packaging, and Labeling",
            "Product Storage and Stability",
            "Preparation",
            "Measures to Minimize Bias: Randomization and Blinding",
            "Study Intervention Compliance",
            "Concomitant Therapy",
            "Rescue Medicine"
        ]
    },
    {
        "id": 7,
        "title": "STUDY INTERVENTION DISCONTINUATION AND PARTICIPANT DISCONTINUATION/WITHDRAWAL",
        "subsections": [
            "Discontinuation of Study Intervention",
            "Participant Discontinuation/Withdrawal from the Study",
            "Lost to Follow-Up"
        ]
    },
    {
        "id": 8,
        "title": "STUDY ASSESSMENTS AND PROCEDURES",
        "subsections": [
            "Efficacy Assessments",
            "Safety and Other Assessments",
            "Adverse Events and Serious Adverse Events",
            "Definition of Adverse Events (AE)",
            "Definition of Serious Adverse Events (SAE)",
            "Classification of an Adverse Event",
            "Time Period and Frequency for Event Assessment and Follow-Up",
            "Adverse Event Reporting",
            "Serious Adverse Event Reporting",
            "Reporting Events to Participants",
            "Events of Special Interest",
            "Reporting of Pregnancy",
            "Unanticipated Problems",
            "Definition of Unanticipated Problems (UP)",
            "Unanticipated Problem Reporting",
            "Reporting Unanticipated Problems to Participants"
        ]
    },
    {
        "id": 9,
        "title": "STATISTICAL CONSIDERATIONS",
        "subsections": [
            "Statistical Hypotheses",
            "Sample Size Determination",
            "Populations for Analyses",
            "Statistical Analyses",
            "General Approach",
            "Analysis of the Primary Efficacy Endpoint(s)",
            "Analysis of the Secondary Endpoint(s)",
            "Safety Analyses",
            "Baseline Descriptive Statistics",
            "Planned Interim Analyses",
            "Sub-Group Analyses",
            "Tabulation of Individual participant Data",
            "Exploratory Analyses"
        ]
    },
    {
        "id": 10,
        "title": "SUPPORTING DOCUMENTATION AND OPERATIONAL CONSIDERATIONS",
        "subsections": [
            "Regulatory, Ethical, and Study Oversight Considerations",
            "Informed Consent Process",
            "Study Discontinuation and Closure",
            "Confidentiality and Privacy",
            "Future Use of Stored Specimens and Data",
            "Key Roles and Study Governance",
            "Safety Oversight",
            "Clinical Monitoring",
            "Quality Assurance and Quality Control",
            "Data Handling and Record Keeping",
            "Protocol Deviations",
            "Publication and Data Sharing Policy",
            "Conflict of Interest Policy",
            "Additional Considerations",
            "Abbreviations",
            "Protocol Amendment History"
        ]
    },
    {
        "id": 11,
        "title": "REFERENCES",
        "subsections": []
    }
]

# ============================================================================
# WORD DOCUMENT GENERATION
# ============================================================================

def get_image_from_data_url(data_url):
    """Convert base64 data URL to a bytes stream for docx"""
    if not data_url:
        return None
    if isinstance(data_url, str) and data_url.startswith('data:image/'):
        try:
            # Format: 'data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAA...'
            header, encoded = data_url.split(",", 1)
            data = base64.b64decode(encoded)
            return io.BytesIO(data)
        except Exception:
            return None
    return None

def set_custom_styles(doc):
    """Set up custom styles matching the template formatting"""
    styles = doc.styles

    # Protocol Title Style
    if 'ProtocolTitle' not in styles:
        title_style = styles.add_style('ProtocolTitle', WD_STYLE_TYPE.PARAGRAPH)
    else:
        title_style = styles['ProtocolTitle']
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(16)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(24)
    title_style.paragraph_format.line_spacing = 1.0

    # Heading 1 Style (Section titles)
    if 'Heading 1' not in styles:
        h1_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    else:
        h1_style = styles['Heading 1']
    h1_style.font.name = 'Calibri'
    h1_style.font.size = Pt(14)
    h1_style.font.bold = True
    h1_style.paragraph_format.space_before = Pt(12)
    h1_style.paragraph_format.space_after = Pt(6)
    h1_style.paragraph_format.keep_with_next = True

    # Heading 2 Style (Subsection titles)
    if 'Heading 2' not in styles:
        h2_style = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    else:
        h2_style = styles['Heading 2']
    h2_style.font.name = 'Calibri'
    h2_style.font.size = Pt(12)
    h2_style.font.bold = True
    h2_style.paragraph_format.space_before = Pt(10)
    h2_style.paragraph_format.space_after = Pt(4)
    h2_style.paragraph_format.left_indent = Pt(0)

    # Heading 3 Style (Sub-subsection titles)
    if 'Heading 3' not in styles:
        h3_style = styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
    else:
        h3_style = styles['Heading 3']
    h3_style.font.name = 'Calibri'
    h3_style.font.size = Pt(11)
    h3_style.font.bold = True
    h3_style.paragraph_format.space_before = Pt(8)
    h3_style.paragraph_format.space_after = Pt(2)
    h3_style.paragraph_format.left_indent = Pt(0) # Reset indent for H3

    # TOC Heading Style (Not included in TOC)
    if 'TOCHeading' not in styles:
        toc_h_style = styles.add_style('TOCHeading', WD_STYLE_TYPE.PARAGRAPH)
    else:
        toc_h_style = styles['TOCHeading']
    toc_h_style.font.name = 'Calibri'
    toc_h_style.font.size = Pt(14)
    toc_h_style.font.bold = True
    toc_h_style.paragraph_format.space_before = Pt(12)
    toc_h_style.paragraph_format.space_after = Pt(6)

    # Normal Text Style
    if 'NormalText' not in styles:
        normal_style = styles.add_style('NormalText', WD_STYLE_TYPE.PARAGRAPH)
    else:
        normal_style = styles['NormalText']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.first_line_indent = Pt(36)

    # Bullet List Style
    if 'BulletList' not in styles:
        bullet_style = styles.add_style('BulletList', WD_STYLE_TYPE.PARAGRAPH)
    else:
        bullet_style = styles['BulletList']
    bullet_style.font.name = 'Calibri'
    bullet_style.font.size = Pt(11)
    bullet_style.paragraph_format.left_indent = Pt(72)
    bullet_style.paragraph_format.first_line_indent = Pt(-18)
    bullet_style.paragraph_format.space_after = Pt(3)

    # Table Text Style
    if 'TableText' not in styles:
        table_style = styles.add_style('TableText', WD_STYLE_TYPE.PARAGRAPH)
    else:
        table_style = styles['TableText']
    table_style.font.name = 'Calibri'
    table_style.font.size = Pt(9)
    table_style.paragraph_format.space_after = Pt(2)

    return styles

def add_custom_header_footer(doc, protocol_data):
    """Add custom header and footer matching the template"""
    # Format date as dd mmm yyyy (e.g., 07 Apr 2017)
    raw_date = protocol_data.get('protocol_date', datetime.now().strftime('%Y-%m-%d'))
    try:
        dt_obj = datetime.strptime(raw_date, '%Y-%m-%d')
        formatted_date = dt_obj.strftime('%d %b %Y')
    except ValueError:
        formatted_date = raw_date

    version_str = protocol_data.get('version_number', '1.0')
    title_str = protocol_data.get('protocol_title', 'Clinical Trial Protocol')

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.header_distance = Cm(1.27)
        section.footer_distance = Cm(1.27)

        # ========== HEADER ==========
        header = section.header
        # clear existing content if any (default is usually one empty paragraph)
        for p in header.paragraphs:
            p.text = ""
            
        # Use a table for Left/Right alignment
        # 1 Row, 2 Columns
        table = header.add_table(rows=1, cols=2, width=Cm(16)) # Width approx page width - margins
        table.autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Remove table borders
        tbl_borders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tbl_borders.append(border)
        table._tbl.tblPr.append(tbl_borders)

        # distinct columns ? allow autofit usually works best for text
        
        # Left Cell: Protocol Title
        cell_left = table.cell(0, 0)
        p_left = cell_left.paragraphs[0]
        p_left.text = title_str
        p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_left.style = doc.styles['Normal']
        # Optional: Add Protocol Number if needed, user only mentioned Title for now.
        
        # Right Cell: Version \n Date
        cell_right = table.cell(0, 1)
        p_right = cell_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_right.style = doc.styles['Normal']
        
        run_v = p_right.add_run(f"Version {version_str}")
        run_v.add_break()
        run_d = p_right.add_run(f"{formatted_date}")

        # ========== FOOTER ==========
        footer = section.footer
        # Clear existing
        for p in footer.paragraphs:
            p.text = ""

        # Line 1: Protocol Title - Version - Date
        # User requested: "protocol title-version dd mmm yyyy"
        footer_text = f"{title_str} – Version {version_str} {formatted_date}"
        p_footer_1 = footer.add_paragraph(footer_text)
        p_footer_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_footer_1.style = doc.styles['Normal']
        p_footer_1.paragraph_format.space_after = Pt(0)

        # Line 2: Page Number
        p_footer_2 = footer.add_paragraph()
        p_footer_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_footer_2.style = doc.styles['Normal']
        
        # Add page number field
        run_page = p_footer_2.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run_page._r.append(fldChar1)
        run_page._r.append(instrText)
        run_page._r.append(fldChar2)

def create_protocol_table(doc, protocol_data):
    """Create the protocol information centered on title page (non-table layout)"""
    # Main Title (centered, bold, larger)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(protocol_data.get('protocol_title', 'Clinical Trial Protocol'))
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph().add_run().add_break()
    
    # Detail items
    details = [
        ("Protocol Number:", protocol_data.get('protocol_number', '')),
        ("National Clinical Trial (NCT) Identified Number:", protocol_data.get('nct_number', '')),
        ("Principal Investigator:", protocol_data.get('principal_investigator', '')),
        ("IND/IDE Sponsor:", protocol_data.get('sponsor', '')),
        ("Funded by:", protocol_data.get('funded_by', '')),
        ("Version Number:", protocol_data.get('version_number', '')),
        ("Date:", protocol_data.get('protocol_date', datetime.now().strftime('%d %B %Y')))
    ]
    
    for label, value in details:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.15
        
        run_label = p.add_run(f"{label} ")
        run_label.bold = True
        run_label.font.size = Pt(12)
        
        run_value = p.add_run(str(value))
        run_value.font.size = Pt(12)

def create_synopsis_section(doc, protocol_data):
    """Create the synopsis section with user inputs"""
    doc.add_heading('1.1 Synopsis', level=2)
    
    # Try new format first (synopsis_data)
    s_data = protocol_data.get('synopsis_data')
    if s_data and s_data.get('overview', {}).get('title'):
        ov = s_data['overview']
        add_labeled_paragraph(doc, "Title:", ov.get('title', ''))
        add_labeled_paragraph(doc, "Coordinating Investigator:", ov.get('coordinating_investigator', ''))
        add_labeled_paragraph(doc, "Planned duration:", ov.get('planned_period', ''))
        
        # Objectives and Endpoints from synopsis_data
        obj = s_data.get('objectives', {})
        if obj.get('primary'):
            doc.add_heading('Primary Objectives', level=3)
            for p_text in obj['primary']:
                p = doc.add_paragraph(style='BulletList')
                add_html_to_paragraph(p, p_text)
        
        # Adding missing Primary Endpoints display match to user's image
        end = s_data.get('endpoints', {})
        if end.get('primary'):
            doc.add_heading('Primary Endpoints', level=3)
            for p_text in end['primary']:
                p = doc.add_paragraph(style='BulletList')
                add_html_to_paragraph(p, p_text)
                
        if s_data.get('num_patients'):
            add_labeled_paragraph(doc, "Number of Patients:", s_data['num_patients'])
            
        return

    # Fallback to legacy synopsis format
    synopsis_data = protocol_data.get('synopsis', {})
    if synopsis_data:
        add_labeled_paragraph(doc, "Title:", synopsis_data.get('Title', ''))
        add_labeled_paragraph(doc, "Study Description:", synopsis_data.get('Study Description', ''))
        add_labeled_paragraph(doc, "Objectives:", synopsis_data.get('Objectives', ''))
        add_labeled_paragraph(doc, "Endpoints:", synopsis_data.get('Endpoints', ''))
        add_labeled_paragraph(doc, "Study Population:", synopsis_data.get('Study Population', ''))
        add_labeled_paragraph(doc, "Phase:", synopsis_data.get('Phase', ''))
        add_labeled_paragraph(doc, "Description of Sites/Facilities:", synopsis_data.get('Sites/Facilities', ''))
        add_labeled_paragraph(doc, "Description of Study Intervention:", synopsis_data.get('Study Intervention Description', ''))
        add_labeled_paragraph(doc, "Study Duration:", synopsis_data.get('Study Duration', ''))
        add_labeled_paragraph(doc, "Participant Duration:", synopsis_data.get('Participant Duration', ''))

def create_approval_section(doc, approval_data):
    """Create the Protocol Approval & Agreement section"""
    if not approval_data or not approval_data.get('details'):
        return

    doc.add_heading('PROTOCOL APPROVAL & AGREEMENT', level=1)
    
    details = approval_data['details']
    
    # Details Table
    table_data = [
        ("Protocol Name:", details.get('protocol_name', '')),
        ("Protocol Number:", details.get('protocol_number', '')),
        ("IMP:", details.get('imp', '')),
        ("Indication:", details.get('indication', '')),
        ("Clinical Phase:", details.get('clinical_phase', '')),
        ("Investigators:", details.get('investigators', '')),
        ("Coordinating Investigator:", details.get('coordinating_investigator', '')),
        ("Expert Committee:", details.get('expert_committee', ''))
    ]
    
    table = doc.add_table(rows=len(table_data), cols=2)
    table.style = 'Table Grid'
    for i, (label, value) in enumerate(table_data):
        table.cell(i, 0).text = label
        table.cell(i, 0).paragraphs[0].runs[0].bold = True
        table.cell(i, 1).text = value
        
    doc.add_paragraph().add_run().add_break()
    
    if details.get('gcp_statement'):
        doc.add_heading('GCP Statement', level=2)
        p = doc.add_paragraph(style='NormalText')
        add_html_to_paragraph(p, details['gcp_statement'])
        
    if details.get('approval_statement'):
        doc.add_heading('Approval Statement', level=2)
        p = doc.add_paragraph(style='NormalText')
        add_html_to_paragraph(p, details['approval_statement'])

    # Sponsor Reps
    if approval_data.get('sponsor_reps'):
        doc.add_heading('Sponsor Representatives', level=2)
        for rep in approval_data['sponsor_reps']:
            p_name = doc.add_paragraph()
            p_name.add_run("Name: ").bold = True
            add_html_to_paragraph(p_name, rep.get('name', ''))
            
            p_title = doc.add_paragraph()
            p_title.add_run("Title: ").bold = True
            add_html_to_paragraph(p_title, rep.get('title', ''))
            
            p_org = doc.add_paragraph()
            p_org.add_run("Organization: ").bold = True
            add_html_to_paragraph(p_org, rep.get('organization', ''))
            
            p_date = doc.add_paragraph()
            p_date.add_run("Date: ").bold = True
            add_html_to_paragraph(p_date, rep.get('date', ''))
            
            # --- DIGITAL SIGNATURE ---
            sig_url = rep.get('signature')
            signature_embedded = False
            
            if sig_url:
                # Try as Base64 first
                sig_stream = get_image_from_data_url(sig_url)
                if sig_stream:
                    try:
                        p_sig_label = doc.add_paragraph()
                        p_sig_label.add_run("Signature: ").bold = True
                        doc.add_picture(sig_stream, width=Inches(1.5))
                        signature_embedded = True
                    except Exception as e:
                        p_sig_err = doc.add_paragraph()
                        p_sig_err.add_run(f"Signature Error: {str(e)}").bold = True
                else:
                    # Try as local path
                    img_path = sig_url.lstrip('/')
                    if os.path.exists(img_path):
                        try:
                            p_sig_label = doc.add_paragraph()
                            p_sig_label.add_run("Signature: ").bold = True
                            doc.add_picture(img_path, width=Inches(1.5))
                            signature_embedded = True
                        except Exception as e:
                            p_sig_err = doc.add_paragraph()
                            p_sig_err.add_run(f"Signature Error: {str(e)}").bold = True
            
            if not signature_embedded:
                # Fallback Signature line
                p_sig = doc.add_paragraph()
                p_sig.add_run("Signature: ").bold = True
                p_sig.add_run("____________________")
            
            doc.add_paragraph("-" * 20)

    # Investigator Agreement
    agree = approval_data.get('investigator_agreement')
    if agree:
        doc.add_heading('Investigator Agreement', level=2)
        doc.add_paragraph(strip_html_tags(agree.get('description', '')), style='NormalText')
        
        p_inv = doc.add_paragraph()
        p_inv.add_run("Investigator Name: ").bold = True
        add_html_to_paragraph(p_inv, agree.get('name', ''))
        
        p_title = doc.add_paragraph()
        p_title.add_run("Title: ").bold = True
        p_title.add_run(strip_html_tags(agree.get('title', '')))
        
        p_facility = doc.add_paragraph()
        p_facility.add_run("Facility: ").bold = True
        p_facility.add_run(strip_html_tags(agree.get('facility', '')))
        
        p_date = doc.add_paragraph()
        p_date.add_run("Date: ").bold = True
        p_date.add_run(strip_html_tags(agree.get('date', '')))

        # --- DIGITAL SIGNATURE ---
        sig_url = agree.get('signature')
        signature_embedded = False
        
        if sig_url:
            sig_stream = get_image_from_data_url(sig_url)
            if sig_stream:
                try:
                    p_sig_label = doc.add_paragraph()
                    p_sig_label.add_run("Signature: ").bold = True
                    doc.add_picture(sig_stream, width=Inches(1.5))
                    signature_embedded = True
                except Exception as e:
                    p_sig_err = doc.add_paragraph()
                    p_sig_err.add_run(f"Signature Error: {str(e)}").bold = True
            else:
                img_path = sig_url.lstrip('/')
                if os.path.exists(img_path):
                    try:
                        p_sig_label = doc.add_paragraph()
                        p_sig_label.add_run("Signature: ").bold = True
                        doc.add_picture(img_path, width=Inches(1.5))
                        signature_embedded = True
                    except Exception as e:
                        p_sig_err = doc.add_paragraph()
                        p_sig_err.add_run(f"Signature Error: {str(e)}").bold = True
        
        if not signature_embedded:
            p_sig = doc.add_paragraph()
            p_sig.add_run("Signature: ").bold = True
            p_sig.add_run("____________________")

def create_soa_table(doc, soa_data):
    """Create the Schedule of Activities table"""
    if not soa_data:
        return

    doc.add_heading('1.3 Schedule of Activities (SoA)', level=2)
    
    # Render Image if exists
    soa_image = soa_data.get('image')
    if soa_image and soa_image.get('url'):
        ip = soa_image.get('url').lstrip('/')
        if os.path.exists(ip):
            try:
                doc.add_picture(ip, width=Inches(6.0))
            except Exception:
                pass
        if soa_image.get('caption'):
            doc.add_paragraph().add_run(str(soa_image.get('caption'))).bold = True
        if soa_image.get('description'):
            doc.add_paragraph(str(soa_image.get('description')))
        doc.add_paragraph() # Add some spacing

    if not soa_data.get('table'):
        return

    headers = soa_data['table'].get('headers', [])
    rows = soa_data['table'].get('rows', [])

    if isinstance(rows, dict):
        table = doc.add_table(rows=1, cols=len(headers) + 1)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Procedures"
        for i, header in enumerate(headers):
            if i + 1 < len(hdr_cells):
                hdr_cells[i+1].text = header
        
        for procedure, row in rows.items():
            cells = table.add_row().cells
            cells[0].text = procedure
            for i, value in enumerate(row):
                if i + 1 < len(cells):
                    cells[i+1].text = "X" if value else ""

    elif isinstance(rows, list):
        if not headers:
            return
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            if i < len(hdr_cells):
                 hdr_cells[i].text = str(header)
        
        for row in rows:
            cells = table.add_row().cells
            for i, val in enumerate(row):
                if i < len(cells):
                    cells[i].text = str(val)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def create_objectives_table(doc, objectives_data):
    """Create the objectives and endpoints table (Legacy support)"""
    if not objectives_data:
        return

    table = doc.add_table(rows=len(objectives_data) + 1, cols=3)
    table.style = 'Table Grid'

    for row in table.rows:
        row.cells[0].width = Cm(3)
        row.cells[1].width = Cm(5)
        row.cells[2].width = Cm(6)

    header_cells = table.rows[0].cells
    header_cells[0].text = "Type"
    header_cells[1].text = "Objectives"
    header_cells[2].text = "Endpoints and Justification"

    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].style = 'TableText'

    for i, row_data in enumerate(objectives_data, start=1):
        cells = table.rows[i].cells
        cells[0].text = row_data.get('Type', '')
        cells[1].text = row_data.get('Objective', '')
        cells[2].text = f"{row_data.get('Endpoint', '')}\n\nJustification: {row_data.get('Justification', '')}"

        for cell in cells:
            cell.paragraphs[0].style = 'TableText'
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

def create_dynamic_table_word(doc, table_data):
    """Create a dynamic table based on headers and rows"""
    if not table_data or not table_data.get('headers'):
        return

    headers = table_data.get('headers', [])
    rows = table_data.get('rows', [])
    
    if not headers:
        return

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = str(header)
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].style = 'TableText'

    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        # row_data might be a list or dict, depending on frontend sending. 
        # Assuming list of values matching headers length
        # Or check if it's a dict and we map by header index?
        # The prompt says "input values in the table", usually implies a grid.
        # Let's assume list of strings.
        
        # Safe extraction
        current_row_len = len(row_data) if isinstance(row_data, list) else 0
        
        for i in range(len(headers)):
            cell = row_cells[i]
            if i < current_row_len:
                cell.text = str(row_data[i])
            else:
                cell.text = ""
            cell.paragraphs[0].style = 'TableText'

def add_images_to_doc(doc, images_list):
    """Helper to add a list of images (dict with url, caption, description) to the doc"""
    if not images_list:
        return
        
    for img in images_list:
        if not img.get('url'):
            continue
            
        img_path = img['url'].lstrip('/')
        if os.path.exists(img_path):
            try:
                # Add picture with fixed width
                doc.add_picture(img_path, width=Cm(15))
            except Exception as e:
                doc.add_paragraph(f"[Error loading image: {str(e)}]", style='NormalText')
        
        if img.get('caption'):
            caption = doc.add_paragraph(style='NormalText')
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.add_run(img['caption']).bold = True
            
        if img.get('description'):
            desc = doc.add_paragraph(img['description'], style='NormalText')
            desc.paragraph_format.space_after = Pt(12)

def create_abbreviations_table(doc, abbreviations_data):
    """Create the abbreviations table"""
    if not abbreviations_data:
        return

    table = doc.add_table(rows=len(abbreviations_data) + 1, cols=2)
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    header_cells[0].text = "Abbreviation"
    header_cells[1].text = "Full Form"

    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].style = 'TableText'

    for i, abbr_data in enumerate(abbreviations_data, start=1):
        cells = table.rows[i].cells
        cells[0].text = abbr_data.get('Abbreviation', '')
        cells[1].text = abbr_data.get('Full Form', '')

        for cell in cells:
            cell.paragraphs[0].style = 'TableText'

def create_amendment_table(doc, amendment_data):
    """Create the protocol amendment history table"""
    if not amendment_data:
        return

    table = doc.add_table(rows=len(amendment_data) + 1, cols=4)
    table.style = 'Table Grid'

    for row in table.rows:
        row.cells[0].width = Cm(2)
        row.cells[1].width = Cm(2)
        row.cells[2].width = Cm(6)
        row.cells[3].width = Cm(6)

    headers = ["Version", "Date", "Description of Change", "Brief Rationale"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].style = 'TableText'

    for i, amend in enumerate(amendment_data, start=1):
        cells = table.rows[i].cells
        cells[0].text = amend.get('Version', '')
        cells[1].text = amend.get('Date', '')
        cells[2].text = amend.get('Description', '')
        cells[3].text = amend.get('Rationale', '')

        for cell in cells:
            cell.paragraphs[0].style = 'TableText'
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

def add_bullet_points(doc, text):
    """Add text as bullet points"""
    if not text:
        return

    lines = text.split('\n')
    for line in lines:
        if line.strip():
            p = doc.add_paragraph(style='BulletList')
            p.add_run(line.strip())

def create_toc(doc):
    """Add Table of Contents field to document"""
    # Create a new paragraph for the TOC
    paragraph = doc.add_paragraph()
    
    # Run 1: Begin Field
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    fldChar.set(qn('w:dirty'), 'true') # Force update on open
    run._r.append(fldChar)
    
    # Run 2: Field Code
    run = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    
    # Run 3: Separator
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar)
    
    # Run 4: Placeholder Text (Visible before update)
    run = paragraph.add_run("Right-click to update Table of Contents")
    
    # Run 5: End Field
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

def add_images_to_doc(doc, images):
    """Helper to add images with captions and descriptions to the document"""
    if not images:
        return
        
    for img in images:
        if not img.get('url'):
            continue
            
        img_path = img['url'].lstrip('/')
        if os.path.exists(img_path):
            try:
                # Add picture - fixed width
                doc.add_picture(img_path, width=Cm(15))
                
                if img.get('caption'):
                    caption = doc.add_paragraph(style='NormalText')
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.add_run(f"Figure: {img['caption']}").bold = True
                    
                if img.get('description'):
                    desc = doc.add_paragraph(img['description'], style='NormalText')
                    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    desc.paragraph_format.space_after = Pt(12)
            except Exception as e:
                print(f"Error adding image {img_path}: {e}")

def generate_complete_word_document(protocol_data):
    """Generate complete Word document matching the template format"""
    doc = Document()
    set_custom_styles(doc)
    add_custom_header_footer(doc, protocol_data)

    # ========== TITLE PAGE ==========
    create_protocol_table(doc, protocol_data)

    if protocol_data.get('summary_changes'):
        doc.add_heading('Summary of Changes from Previous Version:', level=2)
        doc.add_paragraph(protocol_data.get('summary_changes', ''), style='NormalText')

    doc.add_page_break()

    # ========== TABLE OF CONTENTS ==========
    doc.add_paragraph('Table of Contents', style='TOCHeading')
    
    create_toc(doc)

    doc.add_page_break()

    # ========== STATEMENT OF COMPLIANCE ==========
    doc.add_heading('STATEMENT OF COMPLIANCE', level=1)
    compliance_text = """
    The trial will be carried out in accordance with International Conference on Harmonisation Good Clinical Practice (ICH GCP) and the following:

    • United States (US) Code of Federal Regulations (CFR) applicable to clinical studies (45 CFR Part 46, 21 CFR Part 50, 21 CFR Part 56, 21 CFR Part 312, and/or 21 CFR Part 812)

    • National Institutes of Health (NIH)-funded investigators and clinical trial site staff who are responsible for the conduct, management, or oversight of NIH-funded clinical trials have completed Human Subjects Protection and ICH GCP Training.

    • The protocol, informed consent form(s), recruitment materials, and all participant materials will be submitted to the Institutional Review Board (IRB) for review and approval. Approval of both the protocol and the consent form must be obtained before any participant is enrolled. Any amendment to the protocol will require review and approval by the IRB before the changes are implemented to the study. In addition, all changes to the consent form will be IRB-approved; a determination will be made regarding whether a new consent needs to be obtained from participants who provided consent, using a previously approved consent form.
    """
    doc.add_paragraph(compliance_text, style='NormalText')

    doc.add_page_break()

    # ========== PROTOCOL APPROVAL & AGREEMENT ==========
    if protocol_data.get('approval_data'):
        create_approval_section(doc, protocol_data['approval_data'])
        doc.add_page_break()

    # ========== SECTION 1: PROTOCOL SUMMARY ==========
    doc.add_heading('1 PROTOCOL SUMMARY', level=1)

    if protocol_data.get('synopsis_data') or protocol_data.get('synopsis'):
        create_synopsis_section(doc, protocol_data)

    schema_data = protocol_data.get('schema_data')
    if schema_data and (schema_data.get('image_url') or schema_data.get('images')):
        doc.add_heading('1.2 Schema', level=2)
        
        # Prepare list of images to process
        images_to_process = []
        if schema_data.get('images'):
            images_to_process = schema_data['images']
        elif schema_data.get('image_url'):
            # Legacy fallback
            images_to_process = [{
                'url': schema_data['image_url'],
                'caption': schema_data.get('caption'),
                'description': schema_data.get('description')
            }]
            
        for img in images_to_process:
            if not img.get('url'):
                continue
                
            img_path = img['url'].lstrip('/')
            if os.path.exists(img_path):
                # Add picture
                doc.add_picture(img_path, width=Cm(15))
            
            if img.get('caption'):
                caption = doc.add_paragraph(style='NormalText')
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.add_run(img['caption']).bold = True
                
            if img.get('description'):
                desc = doc.add_paragraph(img['description'], style='NormalText')
                # Add some spacing
                desc.paragraph_format.space_after = Pt(12)

    if protocol_data.get('soa_data', {}).get('table'):
        create_soa_table(doc, protocol_data['soa_data'])

    doc.add_page_break()

    # ========== SECTION 2: INTRODUCTION ==========
    sections = protocol_data.get('sections', {})

    if '2' in sections:
        doc.add_heading('2 INTRODUCTION', level=1)

        subsections = sections['2'].get('subsections', [])
        
        # Helper to safely get content from list of dicts
        def get_sub_content(idx):
            if isinstance(subsections, list) and idx < len(subsections):
                 sub = subsections[idx]
                 if isinstance(sub, dict):
                     return sub.get('content')
                 return str(sub) # Fallback if string
            return None

        # 2.1 Study Rationale (Index 0)
        doc.add_heading('2.1 Study Rationale', level=2)
        content = get_sub_content(0)
        if content:
            p = doc.add_paragraph(style='NormalText')
            add_html_to_paragraph(p, content)

        # 2.2 Background (Index 1)
        doc.add_heading('2.2 Background', level=2)
        content = get_sub_content(1)
        if content:
            p = doc.add_paragraph(style='NormalText')
            add_html_to_paragraph(p, content)

        doc.add_heading('2.3 Risk/Benefit Assessment', level=2)

        # 2.3.1 Known Potential Risks (Index 2)
        doc.add_heading('2.3.1 Known Potential Risks', level=3)
        content = get_sub_content(2)
        if content:
            p = doc.add_paragraph(style='NormalText')
            add_html_to_paragraph(p, content)

        # 2.3.2 Known Potential Benefits (Index 3)
        doc.add_heading('2.3.2 Known Potential Benefits', level=3)
        content = get_sub_content(3)
        if content:
            p = doc.add_paragraph(style='NormalText')
            add_html_to_paragraph(p, content)

        # 2.3.3 Assessment of Potential Risks and Benefits (Index 4)
        doc.add_heading('2.3.3 Assessment of Potential Risks and Benefits', level=3)
        content = get_sub_content(4)
        if content:
            p = doc.add_paragraph(style='NormalText')
            add_html_to_paragraph(p, content)

    doc.add_page_break()

    # ========== SECTION 3: OBJECTIVES AND ENDPOINTS ==========
    doc.add_heading('3 OBJECTIVES AND ENDPOINTS', level=1)

    s3_data = protocol_data.get('section3', {})
    
    # 1. Description
    if s3_data.get('description'):
        doc.add_paragraph(s3_data['description'], style='NormalText')
        
    # 2. Image
    if s3_data.get('image', {}).get('url'):
        img_info = s3_data['image']
        img_path = img_info['url'].lstrip('/')
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Cm(15))
            
            # Caption and Description
            if img_info.get('caption'):
                caption = doc.add_paragraph(style='NormalText')
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.add_run(f"Figure: {img_info['caption']}").bold = True
                
            if img_info.get('description'):
                 desc = doc.add_paragraph(img_info['description'], style='NormalText')
                 desc.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3. Dynamic Table
    if s3_data.get('table', {}).get('headers'):
        create_dynamic_table_word(doc, s3_data['table'])

    # Fallback to old objectives_endpoints if section3 is empty (optional, but good for safety)
    if not s3_data.get('description') and not s3_data.get('table', {}).get('headers') and protocol_data.get('objectives_endpoints'):
        create_objectives_table(doc, protocol_data['objectives_endpoints'])

    doc.add_page_break()

    # ========== SECTIONS 4-11 ==========
    for section_num in range(4, 12):
        section_key = str(section_num)
        if section_key in sections:
            # Determine Main Section Title
            section_title = f"{section_num} SECTION {section_num}"
            # Try to get from template structure first
            template_section = next((s for s in TEMPLATE_STRUCTURE if s['id'] == section_num), None)
            if template_section:
                section_title = f"{section_num} {template_section['title']}"
            
            doc.add_heading(section_title, level=1)

            if sections[section_key].get('main'):
                p = doc.add_paragraph(style='NormalText')
                add_html_to_paragraph(p, sections[section_key]['main'])
                
            # Handle images in Main Section
            if sections[section_key].get('images'):
                add_images_to_doc(doc, sections[section_key]['images'])

            # SPECIAL HANDLING FOR SECTION 10 SUBSECTIONS (Abbreviations \u0026 Amendment History)
            # SPECIAL HANDLING FOR SECTION 10 SUBSECTIONS (Abbreviations & Amendment History)
            if section_num == 10:
                # 10.3 Abbreviations
                doc.add_heading('10.3 Abbreviations', level=2)
                if protocol_data.get('abbreviations'):
                    create_abbreviations_table(doc, protocol_data['abbreviations'])
                elif sections[section_key].get('abbreviations'): # Check if inside section data
                    create_abbreviations_table(doc, sections[section_key]['abbreviations'])

                # 10.4 Protocol Amendment History
                doc.add_heading('10.4 Protocol Amendment History', level=2)
                if protocol_data.get('amendment_history'):
                    create_amendment_table(doc, protocol_data['amendment_history'])
                elif sections[section_key].get('amendment_history'):
                    create_amendment_table(doc, sections[section_key]['amendment_history'])
                
                # We skip the normal subsection loop for 10 if we handled them specially
                # or we can let it run if there are OTHER custom subsections.
                # However, usually 10.1 and 10.2 are in the template too.
            
            if sections[section_key].get('subsections'):
                subsections = sections[section_key]['subsections']
                
                # Handle list of dicts (New Format)
                if isinstance(subsections, list):
                    for i, sub in enumerate(subsections):
                        if isinstance(sub, dict):
                            title = sub.get('title', '')
                            content = sub.get('content', '')
                            # Dynamic Numbering
                            full_title = f"{section_num}.{i+1} {title}"
                            doc.add_heading(full_title, level=2)
                            if content:
                                p = doc.add_paragraph(style='NormalText')
                                add_html_to_paragraph(p, content)
                            
                            # Subsection Images
                            if sub.get('images'):
                                add_images_to_doc(doc, sub['images'])
                                
                            # Custom Table
                            if sub.get('customTable'):
                                create_dynamic_table_word(doc, sub['customTable'])
                                
                        elif isinstance(sub, str):
                            # Handle plain string (title only)
                            full_title = f"{section_num}.{i+1} {sub}"
                            doc.add_heading(full_title, level=2)
    
                # Fallback for old dictionary format (Legacy safety)
                elif isinstance(subsections, dict):
                    sorted_subs = sorted(subsections.items(), key=lambda x: int(x[0]))
                    for sub_idx_str, sub_content in sorted_subs:
                        if sub_content:
                            doc.add_heading(f"{section_num}.{int(sub_idx_str)+1} Subsection", level=2)
                            p = doc.add_paragraph(style='NormalText')
                            add_html_to_paragraph(p, sub_content)

            doc.add_page_break()

        
    # ========== CUSTOM SECTIONS (ID > 11) ==========
    custom_keys = [k for k in sections.keys() if k.isdigit() and int(k) > 11]
    custom_keys.sort(key=lambda x: int(x))
    
    for k in custom_keys:
        sec = sections[k]
        doc.add_heading(f"{k} {sec.get('title', 'Section')}", level=1)
        if sec.get('main'):
             doc.add_paragraph(sec['main'], style='NormalText')
             
        if sec.get('images'):
             add_images_to_doc(doc, sec['images'])
             
        if sec.get('subsections') and isinstance(sec['subsections'], list):
             for i, sub in enumerate(sec['subsections']):
                 if isinstance(sub, dict):
                     doc.add_heading(f"{k}.{i+1} {sub.get('title','')}", level=2)
                     if sub.get('content'):
                         doc.add_paragraph(sub['content'], style='NormalText')
                     if sub.get('images'):
                         add_images_to_doc(doc, sub['images'])

    # ========== APPENDICES ==========
    if protocol_data.get('appendices'):
        doc.add_page_break()
        doc.add_heading('APPENDICES', level=1)

        for i, appendix in enumerate(protocol_data['appendices'], 1):
            doc.add_heading(f'Appendix {i}', level=2)
            if appendix.get('title'):
                doc.add_paragraph(appendix['title'], style='NormalText')
            if appendix.get('content'):
                doc.add_paragraph(appendix['content'], style='NormalText')

    # Save document
    output_dir = "generated_docs"
    os.makedirs(output_dir, exist_ok=True)

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    version = protocol_data.get('version_number', 'v1.0').replace('.', '_')
    filename = f"protocol_{version}_{timestamp}.docx"
    filepath = os.path.join(output_dir, filename)
    
    # Update fields on open (prompts user to update TOC)
    element = OxmlElement('w:updateFields')
    element.set(qn('w:val'), 'true')
    doc.settings.element.append(element)

    doc.save(filepath)
    return filepath

# ============================================================================
# PDF DOCUMENT GENERATION
# ============================================================================

def create_pdf_styles():
    """Create PDF styles matching the Word template"""
    styles = getSampleStyleSheet()

    styles.add(ParagraphStyle(
        name='PDFTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=30,
        alignment=1,
        fontName='Helvetica-Bold'
    ))

    styles.add(ParagraphStyle(
        name='PDFHeader1',
        parent=styles['Heading1'],
        fontSize=14,
        spaceBefore=12,
        spaceAfter=6,
        fontName='Helvetica-Bold',
        keepWithNext=True
    ))

    styles.add(ParagraphStyle(
        name='PDFHeader2',
        parent=styles['Heading2'],
        fontSize=12,
        spaceBefore=10,
        spaceAfter=4,
        fontName='Helvetica-Bold',
        leftIndent=0
    ))

    styles.add(ParagraphStyle(
        name='PDFNormal',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=6,
        leading=13,
        firstLineIndent=36
    ))

    styles.add(ParagraphStyle(
        name='PDFTable',
        parent=styles['Normal'],
        fontSize=9,
        spaceAfter=2
    ))

    return styles

def generate_pdf_document(protocol_data):
    """Generate PDF document matching the template format"""
    output_dir = "generated_docs"
    os.makedirs(output_dir, exist_ok=True)

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    version = protocol_data.get('version_number', 'v1.0').replace('.', '_')
    filename = f"protocol_{version}_{timestamp}.pdf"
    filepath = os.path.join(output_dir, filename)

    doc = SimpleDocTemplate(
        filepath,
        pagesize=A4,
        topMargin=1*inch,
        bottomMargin=1*inch,
        leftMargin=1*inch,
        rightMargin=1*inch
    )

    story = []
    styles = create_pdf_styles()

    story.append(Paragraph(protocol_data.get('protocol_title', 'Clinical Trial Protocol'), styles['PDFTitle']))
    story.append(Spacer(1, 0.5*inch))

    info_items = [
        f"<b>Protocol Number:</b> {sanitize_pdf_text(strip_html_tags(protocol_data.get('protocol_number', '')))}",
        f"<b>NCT Number:</b> {sanitize_pdf_text(strip_html_tags(protocol_data.get('nct_number', '')))}",
        f"<b>Principal Investigator:</b> {sanitize_pdf_text(strip_html_tags(protocol_data.get('principal_investigator', '')))}",
        f"<b>Sponsor:</b> {sanitize_pdf_text(strip_html_tags(protocol_data.get('sponsor', '')))}",
        f"<b>Funded By:</b> {sanitize_pdf_text(strip_html_tags(protocol_data.get('funded_by', '')))}",
        f"<b>Version:</b> {sanitize_pdf_text(strip_html_tags(protocol_data.get('version_number', '')))}",
        f"<b>Date:</b> {sanitize_pdf_text(strip_html_tags(protocol_data.get('protocol_date', datetime.now().strftime('%d %B %Y'))))}"
    ]

    for item in info_items:
        story.append(Paragraph(item, styles['PDFNormal']))
        story.append(Spacer(1, 0.1*inch))

    story.append(PageBreak())

    story.append(Paragraph("TABLE OF CONTENTS", styles['PDFHeader1']))
    story.append(Spacer(1, 0.2*inch))

    toc_items = [
        "STATEMENT OF COMPLIANCE",
        "1 PROTOCOL SUMMARY",
        "2 INTRODUCTION",
        "3 OBJECTIVES AND ENDPOINTS",
        "4 STUDY DESIGN",
        "5 STUDY POPULATION",
        "6 STUDY INTERVENTION",
        "7 STUDY INTERVENTION DISCONTINUATION AND PARTICIPANT DISCONTINUATION/WITHDRAWAL",
        "8 STUDY ASSESSMENTS AND PROCEDURES",
        "9 STATISTICAL CONSIDERATIONS",
        "10 SUPPORTING DOCUMENTATION AND OPERATIONAL CONSIDERATIONS",
        "11 REFERENCES",
        "APPENDICES"
    ]

    for item in toc_items:
        story.append(Paragraph(f"• {item}", styles['PDFNormal']))

    story.append(PageBreak())

    story.append(Paragraph("STATEMENT OF COMPLIANCE", styles['PDFHeader1']))

    compliance_text = """
    The trial will be carried out in accordance with International Conference on Harmonisation Good Clinical Practice (ICH GCP) and the following:

    • United States (US) Code of Federal Regulations (CFR) applicable to clinical studies (45 CFR Part 46, 21 CFR Part 50, 21 CFR Part 56, 21 CFR Part 312, and/or 21 CFR Part 812)

    • National Institutes of Health (NIH)-funded investigators and clinical trial site staff who are responsible for the conduct, management, or oversight of NIH-funded clinical trials have completed Human Subjects Protection and ICH GCP Training.

    • The protocol, informed consent form(s), recruitment materials, and all participant materials will be submitted to the Institutional Review Board (IRB) for review and approval.
    """
    story.append(Paragraph(compliance_text, styles['PDFNormal']))

    story.append(PageBreak())

    story.append(Paragraph("1 PROTOCOL SUMMARY", styles['PDFHeader1']))

    if protocol_data.get('synopsis'):
        story.append(Paragraph("1.1 Synopsis", styles['PDFHeader2']))

        table_data = []
        for key, value in protocol_data['synopsis'].items():
            if value:
                table_data.append([key, value])

        if table_data:
            synopsis_table = Table(table_data, colWidths=[2*inch, 4*inch])
            synopsis_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('PADDING', (0, 0), (-1, -1), 6),
            ]))
            story.append(synopsis_table)
    
    # 1.2 Schema
        schema_data = protocol_data.get('schema_data')
        if schema_data and (schema_data.get('image_url') or schema_data.get('images')):
            story.append(Paragraph("1.2 Schema", styles['PDFHeader2']))
            
            # Prepare list of images to process
            images_to_process = []
            if schema_data.get('images'):
                images_to_process = schema_data['images']
            elif schema_data.get('image_url'):
                images_to_process = [{
                    'url': schema_data['image_url'],
                    'caption': schema_data.get('caption'),
                    'description': schema_data.get('description')
                }]

            for img_obj in images_to_process:
                if not img_obj.get('url'):
                    continue

                img_path = img_obj['url'].lstrip('/')
                if os.path.exists(img_path):
                    try:
                        img = Image(img_path, width=6*inch, height=4*inch, kind='proportional')
                        story.append(img)
                    except Exception as e:
                        story.append(Paragraph(f"[Error loading image: {str(e)}]", styles['PDFNormal']))
                
                if img_obj.get('caption'):
                    story.append(Spacer(1, 0.05*inch))
                    story.append(Paragraph(f"<b>{img_obj['caption']}</b>", styles['PDFTable']))

                if img_obj.get('description'):
                    story.append(Paragraph(img_obj['description'], styles['PDFNormal']))
                
                story.append(Spacer(1, 0.2*inch))
             
    # 1.3 SoA
    if protocol_data.get('soa_data', {}).get('table'):
        story.append(Paragraph("1.3 Schedule of Activities (SoA)", styles['PDFHeader2']))
        soa = protocol_data['soa_data']['table']
        # Columns: Procedure + Headers
        headers = ["Procedures"] + soa.get('headers', [])
        data = [headers]
        row_data = soa.get('rows', {})
        for proc, checks in row_data.items():
            row = [proc] + ["X" if c else "" for c in checks]
            data.append(row)
            
        if len(data) > 1:
            # Calculate column widths? lots of columns for SoA.
            # 15 visits + 1 proc = ~16 cols. Page width 6 inch. Each col 0.375 inch? Very small.
            # Might need landscape or very small font.
            # For now, let Table calculate or set minimal.
            soa_table = Table(data)
            soa_table.setStyle(TableStyle([
                ('FONTSIZE', (0,0), (-1,-1), 6),
                ('GRID', (0,0), (-1,-1), 0.5, colors.black),
                ('ALIGN', (1,0), (-1,-1), 'CENTER'),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
            ]))
            story.append(soa_table)

    if protocol_data.get('section3') or protocol_data.get('objectives_endpoints'):
        story.append(Paragraph("3 OBJECTIVES AND ENDPOINTS", styles['PDFHeader1']))
        
        s3_data = protocol_data.get('section3', {})
        
        # Description
        if s3_data.get('description'):
            story.append(Paragraph(s3_data['description'], styles['PDFNormal']))
            story.append(Spacer(1, 0.1*inch))
            
        # Image
        if s3_data.get('image', {}).get('url'):
            img_info = s3_data['image']
            img_path = img_info['url'].lstrip('/')
            if os.path.exists(img_path):
                try:
                    img = Image(img_path, width=6*inch, height=4*inch, kind='proportional')
                    story.append(img)
                    
                    if img_info.get('caption'):
                        story.append(Spacer(1, 0.05*inch))
                        story.append(Paragraph(f"<b>Figure: {img_info['caption']}</b>", styles['PDFTable']))
                        
                    if img_info.get('description'):
                        story.append(Paragraph(img_info['description'], styles['PDFTable']))
                        
                    story.append(Spacer(1, 0.1*inch))
                except Exception as e:
                     story.append(Paragraph(f"[Error loading image: {str(e)}]", styles['PDFNormal']))

        # Dynamic Table
        if s3_data.get('table', {}).get('headers'):
            headers = s3_data['table']['headers']
            rows = s3_data['table']['rows']
            
            if headers:
                table_data = [headers]
                for row in rows:
                     # Ensure row has same length as headers or pad/truncate ?
                     # ReportLab Table needs consistent columns usually or simpler list of lists
                     clean_row = []
                     for i in range(len(headers)):
                         val = row[i] if i < len(row) else ""
                         clean_row.append(str(val))
                     table_data.append(clean_row)
                     
                # Create Table
                # Widths: Auto or distributed? 
                # Let's try to distribute 6 inches among columns
                col_width = 6.5*inch / len(headers) if headers else 1*inch
                
                dyn_table = Table(table_data, colWidths=[col_width]*len(headers))
                dyn_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('PADDING', (0, 0), (-1, -1), 4),
                    ('WORDWRAP', (0, 0), (-1, -1), True),
                ]))
                story.append(dyn_table)

        # Legacy fallback
        elif protocol_data.get('objectives_endpoints'):
             table_data = [["Type", "Objectives", "Endpoints and Justification"]]
             for obj in protocol_data['objectives_endpoints']:
                table_data.append([
                    obj.get('Type', ''),
                    obj.get('Objective', ''),
                    f"{obj.get('Endpoint', '')}\n\nJustification: {obj.get('Justification', '')}"
                ])

             obj_table = Table(table_data, colWidths=[0.8*inch, 2*inch, 3.2*inch])
             obj_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('PADDING', (0, 0), (-1, -1), 4),
                ('WORDWRAP', (0, 0), (-1, -1), True),
             ]))
             story.append(obj_table)

    story.append(PageBreak())

    if protocol_data.get('abbreviations'):
        story.append(Paragraph("10.3 Abbreviations", styles['PDFHeader2']))

        table_data = [["Abbreviation", "Full Form"]]
        for abbr in protocol_data['abbreviations']:
            table_data.append([
                abbr.get('Abbreviation', ''),
                abbr.get('Full Form', '')
            ])

        abbr_table = Table(table_data, colWidths=[1.5*inch, 4.5*inch])
        abbr_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        story.append(abbr_table)

    story.append(PageBreak())

    if protocol_data.get('amendment_history'):
        story.append(Paragraph("10.4 Protocol Amendment History", styles['PDFHeader2']))

        table_data = [["Version", "Date", "Description of Change", "Brief Rationale"]]
        for amend in protocol_data['amendment_history']:
            table_data.append([
                amend.get('Version', ''),
                amend.get('Date', ''),
                amend.get('Description', ''),
                amend.get('Rationale', '')
            ])

        amend_table = Table(table_data, colWidths=[0.8*inch, 1*inch, 2.5*inch, 2.5*inch])
        amend_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('WORDWRAP', (0, 0), (-1, -1), True),
        ]))
        story.append(amend_table)

    doc.build(story)
    return filepath

def generate_word_document(protocol_data):
    """Main function to generate Word document"""
    return generate_complete_word_document(protocol_data)


def generate_interpreted_word_report(protocol_id):
    """
    Generates a specialized Word report containing only the 12 interpreted fields.
    Highlights low-confidence extractions in red.
    """
    from database import execute_query
    
    # 1. Fetch data
    query = "SELECT field_name, field_value, confidence_score FROM protocol_interpretation WHERE protocol_id = %s ORDER BY field_name"
    data = execute_query(query, (protocol_id,), fetch=True)
    
    if not data:
        logger.warning(f"No interpreted data found for protocol_id {protocol_id}")
    
    # 2. Create Document
    doc = Document()
    
    # Header Info
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = f"Protocol Interpretation Report | ID: {protocol_id} | Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    title = doc.add_heading("Protocol Interpretation Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("This report summarizes the core study parameters extracted and interpreted from the protocol source data.")
    
    # 3. Build Content (Non-Tabular)
    for item in data:
        # Add Field Name as a Heading
        doc.add_heading(str(item["field_name"]), level=2)
        
        # Add Field Value as a Paragraph
        val_text = str(item["field_value"])
        p = doc.add_paragraph(style="NormalText")
        run = p.add_run(val_text)
        
        # Check confidence for styling
        conf = float(item["confidence_score"] or 1.0)
        if conf < 1.0:
            run.font.color.rgb = RGBColor(255, 0, 0) # Red
            # Add a small note about confidence
            conf_note = p.add_run(f" (Confidence: {int(conf * 100)}%)")
            conf_note.italic = True
            conf_note.font.size = Pt(9)
        
    doc.add_paragraph("\nNote: Values highlighted in red indicate lower extraction confidence and should be manually verified by a domain expert.")
    
    # Footer
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Confidential - Clinical Trial Protocol Interpretation System"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc
